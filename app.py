# File: app_ultra_pro.py

import streamlit as st
import google.generativeai as genai
from PIL import Image, ImageEnhance, ImageOps, ImageFilter
import pypdfium2 as pdfium
import io, gc, base64, json
from datetime import datetime
from docx import Document
from supabase import create_client
from google.generativeai.types import HarmCategory, HarmBlockThreshold

# ==========================================
# 1. TIZIM VA SEO SOZLAMALARI
# ==========================================
st.set_page_config(
    page_title="Manuscript AI - Global Academic Master",
    page_icon="📜",
    layout="wide",
    initial_sidebar_state="auto"
)

# ==========================================
# SESSION STATE INITIALIZATION
# ==========================================
if "auth" not in st.session_state:
    st.session_state.auth = False
if "u_email" not in st.session_state:
    st.session_state.u_email = ""
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = False
if "theme" not in st.session_state:
    st.session_state.theme = "Antik Oltin"
if "history" not in st.session_state:
    st.session_state.history = []
if "imgs" not in st.session_state:
    st.session_state.imgs = []
if "results" not in st.session_state:
    st.session_state.results = {}
if "chats" not in st.session_state:
    st.session_state.chats = {}
if "compare_mode" not in st.session_state:
    st.session_state.compare_mode = False
if "current_page_index" not in st.session_state:
    st.session_state.current_page_index = 0
if "show_landing" not in st.session_state:
    st.session_state.show_landing = False
if "file_data" not in st.session_state:
    st.session_state.file_data = None
if "is_pdf" not in st.session_state:
    st.session_state.is_pdf = False
if "total_pages" not in st.session_state:
    st.session_state.total_pages = 0

# ==========================================
# THEME DEFINITIONS
# ==========================================
THEMES = {
    "Antik Oltin": {
        "primary": "#0c1421",
        "accent": "#c5a059",
        "accent2": "#d4af37",
        "bg_light": "#f4ecd8",
        "bg_light2": "#fdfaf1",
        "bg_dark": "#1a1a1a",
        "bg_dark2": "#2d2d2d"
    },
    "Moviy Professional": {
        "primary": "#1e3a8a",
        "accent": "#3b82f6",
        "accent2": "#60a5fa",
        "bg_light": "#eff6ff",
        "bg_light2": "#dbeafe",
        "bg_dark": "#1e293b",
        "bg_dark2": "#334155"
    },
    "Yashil Tabiat": {
        "primary": "#065f46",
        "accent": "#10b981",
        "accent2": "#34d399",
        "bg_light": "#ecfdf5",
        "bg_light2": "#d1fae5",
        "bg_dark": "#1a2e1a",
        "bg_dark2": "#2d4a2d"
    }
}

theme = THEMES[st.session_state.theme]
dark = st.session_state.dark_mode

# Dynamic color variables
bg_main = theme["bg_dark"] if dark else theme["bg_light"]
bg_secondary = theme["bg_dark2"] if dark else theme["bg_light2"]
text_primary = "#e0e0e0" if dark else "#1a1a1a"
text_secondary = "#a0a0a0" if dark else "#5d4037"
card_bg = theme["bg_dark2"] if dark else "#ffffff"

# ==========================================
# DYNAMIC CSS - ULTRA PROFESSIONAL
# ==========================================
st.markdown(f"""
    <style>
    /* === GOOGLE FONTS === */
    @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;500;600;700&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
    
    /* === CSS VARIABLES === */
    :root {{
        --primary: {theme['primary']};
        --accent: {theme['accent']};
        --accent2: {theme['accent2']};
        --bg-main: {bg_main};
        --bg-secondary: {bg_secondary};
        --text-primary: {text_primary};
        --text-secondary: {text_secondary};
        --card-bg: {card_bg};
        --glass-bg: rgba(255,255,255,0.05);
        --glass-border: rgba(255,255,255,0.1);
        --shadow-sm: 0 2px 8px rgba(0,0,0,0.08);
        --shadow-md: 0 8px 24px rgba(0,0,0,0.12);
        --shadow-lg: 0 16px 48px rgba(0,0,0,0.16);
        --shadow-glow: 0 0 40px rgba(197,160,89,0.3);
        --transition-fast: 0.15s cubic-bezier(0.4, 0, 0.2, 1);
        --transition-normal: 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        --transition-slow: 0.5s cubic-bezier(0.4, 0, 0.2, 1);
        --blur-amount: 20px;
    }}
    
    /* === SYSTEM OVERRIDES === */
    footer {{visibility: hidden !important;}}
    .stAppDeployButton {{display:none !important;}}
    #stDecoration {{display:none !important;}}
    
    /* === SCROLLBAR STYLING === */
    ::-webkit-scrollbar {{
        width: 8px;
        height: 8px;
    }}
    
    ::-webkit-scrollbar-track {{
        background: {bg_secondary};
        border-radius: 4px;
    }}
    
    ::-webkit-scrollbar-thumb {{
        background: linear-gradient(180deg, {theme['accent']}, {theme['accent2']});
        border-radius: 4px;
        transition: var(--transition-normal);
    }}
    
    ::-webkit-scrollbar-thumb:hover {{
        background: {theme['accent2']};
    }}

    header[data-testid='stHeader'] {{
        background: rgba(0,0,0,0) !important;
        visibility: visible !important;
        backdrop-filter: blur(10px);
    }}

    button[data-testid='stSidebarCollapseButton'] {{
        background: linear-gradient(135deg, {theme['primary']} 0%, rgba(12,20,33,0.9) 100%) !important;
        color: {theme['accent']} !important;
        border: 1px solid rgba(197,160,89,0.3) !important;
        position: fixed !important;
        z-index: 1000001 !important;
        transition: var(--transition-normal) !important;
        backdrop-filter: blur(10px);
        box-shadow: var(--shadow-md);
    }}
    
    button[data-testid='stSidebarCollapseButton']:hover {{
        background: linear-gradient(135deg, {theme['accent']} 0%, {theme['accent2']} 100%) !important;
        color: {theme['primary']} !important;
        transform: scale(1.08) rotate(3deg);
        box-shadow: var(--shadow-glow);
    }}

    /* === MAIN LAYOUT === */
    .main {{ 
        background: linear-gradient(135deg, {bg_main} 0%, {bg_secondary} 50%, {bg_main} 100%) !important;
        color: {text_primary} !important;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
        padding: 2rem 1rem;
        min-height: 100vh;
        position: relative;
    }}
    
    /* Subtle animated background pattern */
    .main::before {{
        content: '';
        position: fixed;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-image: 
            radial-gradient(circle at 20% 80%, rgba(197,160,89,0.03) 0%, transparent 50%),
            radial-gradient(circle at 80% 20%, rgba(197,160,89,0.03) 0%, transparent 50%),
            radial-gradient(circle at 40% 40%, rgba(197,160,89,0.02) 0%, transparent 30%);
        pointer-events: none;
        z-index: 0;
    }}
    
    /* === TYPOGRAPHY === */
    h1, h2, h3, h4 {{ 
        color: {theme['primary'] if not dark else theme['accent']} !important;
        font-family: 'Playfair Display', Georgia, serif;
        font-weight: 600;
        letter-spacing: -0.02em;
        text-align: center;
        padding-bottom: 16px;
        margin-bottom: 24px;
        position: relative;
    }}
    
    h1::after, h2::after {{
        content: '';
        position: absolute;
        bottom: 0;
        left: 50%;
        transform: translateX(-50%);
        width: 80px;
        height: 3px;
        background: linear-gradient(90deg, transparent, {theme['accent']}, transparent);
        border-radius: 2px;
    }}
    
    h1 {{
        font-size: clamp(2rem, 5vw, 3rem) !important;
        margin-top: 0;
        text-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }}
    
    h2 {{
        font-size: clamp(1.5rem, 3.5vw, 2.2rem) !important;
    }}
    
    h3 {{
        font-size: clamp(1.2rem, 3vw, 1.6rem) !important;
        border-bottom: none;
    }}
    
    h4 {{
        font-size: clamp(1.1rem, 2.5vw, 1.4rem) !important;
        border-bottom: none;
    }}
    
    p, li, span {{
        font-family: 'Inter', sans-serif;
        line-height: 1.7;
    }}

    /* === BUTTONS - GLASSMORPHISM === */
    .stButton>button {{ 
        background: linear-gradient(135deg, {theme['primary']} 0%, rgba(12,20,33,0.95) 100%) !important;
        color: {theme['accent']} !important;
        font-family: 'Inter', sans-serif !important;
        font-weight: 600 !important;
        letter-spacing: 0.02em !important;
        width: 100% !important;
        padding: 16px 24px !important;
        border: 1px solid rgba(197,160,89,0.3) !important;
        border-radius: 12px !important;
        height: auto !important;
        min-height: 54px !important;
        font-size: 15px !important;
        cursor: pointer !important;
        transition: all var(--transition-normal) !important;
        box-shadow: var(--shadow-md), inset 0 1px 0 rgba(255,255,255,0.05) !important;
        backdrop-filter: blur(10px) !important;
        position: relative !important;
        overflow: hidden !important;
    }}
    
    .stButton>button::before {{
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(197,160,89,0.2), transparent);
        transition: left 0.5s ease;
        pointer-events: none;
    }}
    
    .stButton>button:hover::before {{
        left: 100%;
    }}
    
    .stButton>button:hover {{ 
        transform: translateY(-3px) !important;
        box-shadow: var(--shadow-lg), var(--shadow-glow) !important;
        border-color: {theme['accent']} !important;
        background: linear-gradient(135deg, {theme['primary']} 0%, {theme['accent']}22 100%) !important;
        color: {theme['accent2']} !important;
    }}
    
    .stButton>button:active {{
        transform: translateY(-1px) !important;
        box-shadow: var(--shadow-sm) !important;
    }}
    
    /* Primary CTA Button */
    .stButton>button[kind="primary"] {{
        background: linear-gradient(135deg, {theme['accent']} 0%, {theme['accent2']} 100%) !important;
        color: {theme['primary']} !important;
        border: none !important;
    }}

    /* === RESULT BOX - PREMIUM GLASSMORPHISM === */
    .result-box {{ 
        background: linear-gradient(145deg, rgba(30,42,56,0.95) 0%, rgba(22,32,44,0.98) 100%) !important;
        padding: 32px !important;
        border-radius: 16px !important;
        border: 1px solid rgba(197,160,89,0.25) !important;
        border-left: 4px solid {theme['accent']} !important;
        box-shadow: var(--shadow-lg), inset 0 1px 0 rgba(255,255,255,0.05) !important;
        backdrop-filter: blur(20px) !important;
        color: {text_primary} !important;
        font-size: 16px;
        line-height: 1.85;
        margin-bottom: 24px;
        animation: fadeInUp 0.6s cubic-bezier(0.23, 1, 0.32, 1);
        position: relative;
        overflow: hidden;
        transition: all var(--transition-normal);
    }}
    
    .result-box::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 2px;
        background: linear-gradient(90deg, {theme['accent']}, {theme['accent2']}, {theme['accent']});
        background-size: 200% 100%;
        animation: shimmer 3s ease-in-out infinite;
        pointer-events: none;
    }}
    
    @keyframes shimmer {{
        0%, 100% {{ background-position: 200% 0; }}
        50% {{ background-position: -200% 0; }}
    }}
    
    @keyframes float {{
        0%, 100% {{ transform: translateY(0); }}
        50% {{ transform: translateY(-8px); }}
    }}
    
    .result-box:hover {{
        border-color: rgba(197,160,89,0.4) !important;
        box-shadow: var(--shadow-xl), var(--shadow-glow) !important;
        transform: translateY(-2px);
    }}
    
    /* Result box ichidagi markdown elementlari */
    .result-box h1, .result-box h2, .result-box h3 {{
        color: {theme['accent']} !important;
        font-family: 'Playfair Display', serif !important;
        margin-top: 20px !important;
        margin-bottom: 12px !important;
        padding-bottom: 8px !important;
        border-bottom: 1px solid rgba(197,160,89,0.2) !important;
    }}
    
    .result-box h2 {{
        font-size: 1.3rem !important;
    }}
    
    .result-box h3 {{
        font-size: 1.1rem !important;
    }}
    
    .result-box p {{
        margin-bottom: 12px !important;
        line-height: 1.9 !important;
    }}
    
    .result-box strong, .result-box b {{
        color: {theme['accent2']} !important;
    }}
    
    .result-box hr {{
        border: none !important;
        height: 1px !important;
        background: linear-gradient(90deg, transparent, rgba(197,160,89,0.3), transparent) !important;
        margin: 16px 0 !important;
    }}
    
    .result-box ul, .result-box ol {{
        padding-left: 20px !important;
        margin-bottom: 12px !important;
    }}
    
    .result-box li {{
        margin-bottom: 6px !important;
    }}
    
    @keyframes fadeInUp {{
        from {{
            opacity: 0;
            transform: translateY(30px);
        }}
        to {{
            opacity: 1;
            transform: translateY(0);
        }}
    }}

    /* === FORM INPUTS - PREMIUM === */
    .stTextInput>div>div>input,
    .stTextArea textarea {{
        background: linear-gradient(135deg, {bg_secondary} 0%, rgba(30,42,56,0.95) 100%) !important;
        color: {text_primary} !important;
        border: 1px solid rgba(197,160,89,0.3) !important;
        border-radius: 12px !important;
        padding: 14px 16px !important;
        font-family: 'JetBrains Mono', 'Fira Code', monospace !important;
        transition: all var(--transition-normal) !important;
        font-size: 14px !important;
        box-shadow: var(--shadow-sm), inset 0 2px 4px rgba(0,0,0,0.1) !important;
    }}
    
    .stTextInput>div>div>input:focus,
    .stTextArea textarea:focus {{
        border-color: {theme['accent']} !important;
        box-shadow: 0 0 0 3px rgba(197,160,89,0.15), var(--shadow-md) !important;
        outline: none !important;
        background: linear-gradient(135deg, rgba(30,42,56,1) 0%, rgba(22,32,44,1) 100%) !important;
    }}
    
    .stTextInput>div>div>input::placeholder,
    .stTextArea textarea::placeholder {{
        color: rgba(253,250,241,0.4) !important;
        font-style: italic;
    }}

    /* === CHAT BUBBLES - PREMIUM === */
    .chat-user {{
        background: linear-gradient(135deg, rgba(30,42,56,0.95) 0%, rgba(22,32,44,0.9) 100%);
        color: {text_primary} !important;
        padding: 18px 22px;
        border-radius: 20px 20px 6px 20px;
        border: 1px solid rgba(12,20,33,0.3);
        border-left: 4px solid {theme['primary']};
        margin-bottom: 14px;
        box-shadow: var(--shadow-md);
        animation: slideInLeft 0.4s cubic-bezier(0.23, 1, 0.32, 1);
        backdrop-filter: blur(10px);
        position: relative;
    }}
    
    .chat-user::after {{
        content: '👤';
        position: absolute;
        bottom: -8px;
        left: 16px;
        font-size: 14px;
    }}
    
    .chat-ai {{
        background: linear-gradient(135deg, rgba(22,32,44,0.95) 0%, rgba(30,42,56,0.9) 100%);
        color: {text_primary} !important;
        padding: 18px 22px;
        border-radius: 20px 20px 20px 6px;
        border: 1px solid rgba(197,160,89,0.2);
        border-left: 4px solid {theme['accent']};
        margin-bottom: 18px;
        box-shadow: var(--shadow-md), 0 0 20px rgba(197,160,89,0.08);
        animation: slideInRight 0.4s cubic-bezier(0.23, 1, 0.32, 1);
        backdrop-filter: blur(10px);
        position: relative;
    }}
    
    .chat-ai::after {{
        content: '🤖';
        position: absolute;
        bottom: -8px;
        right: 16px;
        font-size: 14px;
    }}
    
    @keyframes slideInLeft {{
        from {{ opacity: 0; transform: translateX(-20px); }}
        to {{ opacity: 1; transform: translateX(0); }}
    }}
    
    @keyframes slideInRight {{
        from {{ opacity: 0; transform: translateX(20px); }}
        to {{ opacity: 1; transform: translateX(0); }}
    }}

    /* === SIDEBAR - PREMIUM === */
    section[data-testid='stSidebar'] {{
        background: linear-gradient(180deg, {theme['primary']} 0%, #0a0f18 100%) !important;
        border-right: 1px solid rgba(197,160,89,0.3) !important;
        box-shadow: 4px 0 20px rgba(0,0,0,0.3) !important;
    }}
    
    section[data-testid='stSidebar']::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background-image: url("data:image/svg+xml,%3Csvg width='60' height='60' viewBox='0 0 60 60' xmlns='http://www.w3.org/2000/svg'%3E%3Cg fill='none' fill-rule='evenodd'%3E%3Cg fill='%23c5a059' fill-opacity='0.03'%3E%3Cpath d='M36 34v-4h-2v4h-4v2h4v4h2v-4h4v-2h-4zm0-30V0h-2v4h-4v2h4v4h2V6h4V4h-4zM6 34v-4H4v4H0v2h4v4h2v-4h4v-2H6zM6 4V0H4v4H0v2h4v4h2V6h4V4H6z'/%3E%3C/g%3E%3C/g%3E%3C/svg%3E");
        pointer-events: none;
        z-index: 0;
    }}
    
    section[data-testid='stSidebar'] .stMarkdown {{
        color: #fdfaf1 !important;
        position: relative;
        z-index: 1;
    }}
    
    section[data-testid='stSidebar'] [data-testid='stMetricValue'] {{
        color: {theme['accent']} !important;
        font-size: 28px !important;
        font-weight: 700 !important;
        font-family: 'Playfair Display', serif !important;
        text-shadow: 0 2px 4px rgba(0,0,0,0.2);
    }}
    
    section[data-testid='stSidebar'] [data-testid='stMetricLabel'] {{
        color: rgba(253,250,241,0.7) !important;
        font-family: 'Inter', sans-serif !important;
        font-size: 12px !important;
        text-transform: uppercase !important;
        letter-spacing: 0.1em !important;
    }}

    /* === IMAGE MAGNIFIER - PREMIUM === */
    .magnifier-container {{
        overflow: hidden;
        border: 2px solid rgba(197,160,89,0.4);
        border-radius: 16px;
        cursor: zoom-in;
        background: linear-gradient(135deg, #ffffff 0%, #f8f6f0 100%);
        padding: 12px;
        box-shadow: var(--shadow-lg), inset 0 0 20px rgba(0,0,0,0.03);
        transition: all var(--transition-normal);
        position: relative;
    }}
    
    .magnifier-container::before {{
        content: '🔍';
        position: absolute;
        top: 16px;
        right: 16px;
        font-size: 20px;
        opacity: 0;
        transition: opacity var(--transition-fast);
        z-index: 10;
        background: rgba(255,255,255,0.9);
        padding: 6px 10px;
        border-radius: 8px;
        box-shadow: var(--shadow-sm);
        pointer-events: none;
    }}
    
    .magnifier-container:hover::before {{
        opacity: 1;
    }}
    
    .magnifier-container:hover {{
        box-shadow: var(--shadow-xl), var(--shadow-glow);
        border-color: {theme['accent']};
        transform: translateY(-4px);
    }}
    
    .magnifier-container img {{
        transition: transform 0.5s cubic-bezier(0.25, 0.46, 0.45, 0.94);
        border-radius: 8px;
    }}
    
    .magnifier-container img:hover {{
        transform: scale(1.4);
    }}

    /* === ZOOM MODAL - PREMIUM === */
    .modal {{
        display: none;
        position: fixed;
        z-index: 99999;
        padding-top: 50px;
        left: 0;
        top: 0;
        width: 100%;
        height: 100%;
        overflow: auto;
        background-color: rgba(0,0,0,0.97);
        animation: fadeIn 0.3s cubic-bezier(0.23, 1, 0.32, 1);
        backdrop-filter: blur(10px);
    }}
    
    .modal-content {{
        margin: auto;
        display: block;
        width: 90%;
        max-width: 1400px;
        animation: zoomIn 0.4s cubic-bezier(0.23, 1, 0.32, 1);
        border-radius: 12px;
        box-shadow: 0 25px 80px rgba(0,0,0,0.5);
    }}
    
    @keyframes fadeIn {{
        from {{ opacity: 0; }}
        to {{ opacity: 1; }}
    }}
    
    @keyframes zoomIn {{
        from {{ transform: scale(0.8) translateY(20px); opacity: 0; }}
        to {{ transform: scale(1) translateY(0); opacity: 1; }}
    }}
    
    .modal-close {{
        position: absolute;
        top: 20px;
        right: 40px;
        color: rgba(255,255,255,0.8);
        font-size: 44px;
        font-weight: 300;
        transition: all var(--transition-fast);
        cursor: pointer;
        z-index: 100000;
        width: 50px;
        height: 50px;
        display: flex;
        align-items: center;
        justify-content: center;
        border-radius: 50%;
        background: rgba(255,255,255,0.1);
        backdrop-filter: blur(10px);
    }}
    
    .modal-close:hover,
    .modal-close:focus {{
        color: {theme['accent']};
        background: rgba(197,160,89,0.2);
        transform: rotate(90deg);
    }}

    /* === CITATION BOX - PREMIUM === */
    .citation-box {{
        font-size: 13px;
        color: {text_secondary};
        background: linear-gradient(135deg, rgba(30,42,56,0.6) 0%, rgba(22,32,44,0.8) 100%);
        padding: 18px 22px;
        border-radius: 12px;
        border: 1px dashed rgba(197,160,89,0.4);
        margin-top: 20px;
        font-style: italic;
        box-shadow: var(--shadow-sm);
        position: relative;
        overflow: hidden;
    }}
    
    .citation-box::before {{
        content: '📚';
        position: absolute;
        top: 50%;
        left: -30px;
        transform: translateY(-50%);
        font-size: 60px;
        opacity: 0.05;
    }}

    /* === FILE UPLOADER - PREMIUM === */
    [data-testid='stFileUploader'] {{
        background: linear-gradient(135deg, rgba(30,42,56,0.5) 0%, rgba(22,32,44,0.7) 100%);
        border: 2px dashed rgba(197,160,89,0.4);
        border-radius: 16px;
        padding: 40px;
        text-align: center;
        transition: all var(--transition-normal);
        position: relative;
        overflow: hidden;
    }}
    
    [data-testid='stFileUploader']::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: linear-gradient(45deg, transparent 30%, rgba(197,160,89,0.03) 50%, transparent 70%);
        background-size: 200% 200%;
        animation: uploadShimmer 3s ease-in-out infinite;
        pointer-events: none;
        z-index: 0;
    }}
    
    @keyframes uploadShimmer {{
        0% {{ background-position: 200% 200%; }}
        100% {{ background-position: -200% -200%; }}
    }}
    
    [data-testid='stFileUploader']:hover {{
        border-color: {theme['accent']};
        background: linear-gradient(135deg, rgba(30,42,56,0.7) 0%, rgba(22,32,44,0.9) 100%);
        box-shadow: var(--shadow-lg), var(--shadow-glow);
        transform: translateY(-2px);
    }}
    
    /* === EXPANDER STYLING === */
    .streamlit-expanderHeader {{
        background: linear-gradient(135deg, rgba(30,42,56,0.8) 0%, rgba(22,32,44,0.9) 100%) !important;
        border: 1px solid rgba(197,160,89,0.2) !important;
        border-radius: 12px !important;
        padding: 14px 18px !important;
        font-family: 'Inter', sans-serif !important;
        font-weight: 600 !important;
        color: {theme['accent']} !important;
        transition: all var(--transition-normal) !important;
    }}
    
    .streamlit-expanderHeader:hover {{
        background: linear-gradient(135deg, rgba(30,42,56,0.95) 0%, rgba(22,32,44,1) 100%) !important;
        border-color: {theme['accent']} !important;
        box-shadow: var(--shadow-md) !important;
    }}
    
    .streamlit-expanderContent {{
        background: rgba(22,32,44,0.5) !important;
        border: 1px solid rgba(197,160,89,0.1) !important;
        border-top: none !important;
        border-radius: 0 0 12px 12px !important;
        padding: 16px !important;
    }}
    
    /* === STATUS INDICATOR === */
    [data-testid='stStatusWidget'] {{
        background: linear-gradient(135deg, rgba(30,42,56,0.9) 0%, rgba(22,32,44,0.95) 100%) !important;
        border: 1px solid rgba(197,160,89,0.2) !important;
        border-radius: 12px !important;
        padding: 16px !important;
    }}
    
    /* === SELECTBOX STYLING === */
    [data-testid='stSelectbox'] > div > div {{
        background: linear-gradient(135deg, rgba(30,42,56,0.8) 0%, rgba(22,32,44,0.9) 100%) !important;
        border: 1px solid rgba(197,160,89,0.3) !important;
        border-radius: 10px !important;
        color: {text_primary} !important;
    }}
    
    [data-testid='stSelectbox'] > div > div:hover {{
        border-color: {theme['accent']} !important;
    }}
    
    /* === MULTISELECT STYLING === */
    [data-testid='stMultiSelect'] > div > div {{
        background: linear-gradient(135deg, rgba(30,42,56,0.8) 0%, rgba(22,32,44,0.9) 100%) !important;
        border: 1px solid rgba(197,160,89,0.3) !important;
        border-radius: 10px !important;
    }}
    
    [data-testid='stMultiSelect'] span[data-baseweb='tag'] {{
        background: linear-gradient(135deg, {theme['accent']} 0%, {theme['accent2']} 100%) !important;
        color: {theme['primary']} !important;
        border-radius: 6px !important;
        font-weight: 600 !important;
    }}

    /* Ensure file uploader children are clickable */
    [data-testid='stFileUploader'] > div,
    [data-testid='stFileUploader'] button,
    [data-testid='stFileUploader'] input,
    [data-testid='stFileUploader'] label,
    [data-testid='stFileUploader'] section,
    [data-testid='stFileUploader'] [data-testid] {{
        position: relative !important;
        z-index: 1 !important;
    }}

    /* === DIVIDER - PREMIUM === */
    hr {{
        border: none;
        height: 1px;
        background: linear-gradient(90deg, transparent 0%, rgba(197,160,89,0.3) 20%, {theme['accent']} 50%, rgba(197,160,89,0.3) 80%, transparent 100%);
        margin: 40px 0;
        position: relative;
    }}
    
    hr::after {{
        content: '✦';
        position: absolute;
        left: 50%;
        top: 50%;
        transform: translate(-50%, -50%);
        background: {theme['primary']};
        padding: 0 12px;
        color: {theme['accent']};
        font-size: 14px;
    }}

    /* === EMPTY STATE - PREMIUM === */
    .empty-state {{
        text-align: center;
        padding: 80px 30px;
        background: linear-gradient(145deg, rgba(30,42,56,0.6) 0%, rgba(22,32,44,0.8) 100%);
        border-radius: 20px;
        border: 2px dashed rgba(197,160,89,0.3);
        margin: 50px 0;
        box-shadow: var(--shadow-lg);
        position: relative;
        overflow: hidden;
    }}
    
    .empty-state::before {{
        content: '';
        position: absolute;
        top: -50%;
        left: -50%;
        width: 200%;
        height: 200%;
        background: radial-gradient(circle, rgba(197,160,89,0.05) 0%, transparent 70%);
        animation: pulse 4s ease-in-out infinite;
        pointer-events: none;
    }}
    
    @keyframes pulse {{
        0%, 100% {{ transform: scale(1); opacity: 0.5; }}
        50% {{ transform: scale(1.1); opacity: 0.8; }}
    }}
    
    .empty-state h3 {{
        color: {theme['accent']};
        border: none;
        margin-bottom: 12px;
        position: relative;
        z-index: 1;
    }}
    
    .empty-state p {{
        position: relative;
        z-index: 1;
    }}

    /* === LOGIN CARD - PREMIUM === */
    .login-card {{
        background: linear-gradient(145deg, rgba(30,42,56,0.98) 0%, rgba(22,32,44,0.99) 100%);
        padding: 60px 50px;
        border-radius: 24px;
        box-shadow: var(--shadow-xl), 0 0 60px rgba(197,160,89,0.15);
        border: 1px solid rgba(197,160,89,0.3);
        animation: fadeInUp 0.8s cubic-bezier(0.23, 1, 0.32, 1);
        backdrop-filter: blur(20px);
        position: relative;
        overflow: hidden;
    }}
    
    .login-card::before {{
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: linear-gradient(90deg, {theme['accent']}, {theme['accent2']}, {theme['accent']});
        background-size: 200% 100%;
        animation: shimmer 3s ease-in-out infinite;
        pointer-events: none;
    }}
    
    .login-card::after {{
        content: '';
        position: absolute;
        bottom: 0;
        left: 50%;
        transform: translateX(-50%);
        width: 80%;
        height: 1px;
        background: linear-gradient(90deg, transparent, rgba(197,160,89,0.3), transparent);
    }}
    
    .hero-title {{
        font-size: clamp(2.2rem, 5vw, 3.5rem);
        color: {theme['accent']};
        margin-bottom: 18px;
        font-weight: 700;
        text-align: center;
        font-family: 'Playfair Display', serif;
        letter-spacing: -0.02em;
        text-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }}
    
    .hero-subtitle {{
        font-size: clamp(1rem, 2vw, 1.15rem);
        color: rgba(253,250,241,0.7);
        text-align: center;
        margin-bottom: 40px;
        line-height: 1.7;
        font-family: 'Inter', sans-serif;
        max-width: 400px;
        margin-left: auto;
        margin-right: auto;
    }}

    /* === CREDIT PROGRESS BAR - PREMIUM === */
    .credit-bar-container {{
        background: rgba(0,0,0,0.3);
        border-radius: 12px;
        padding: 5px;
        margin: 18px 0;
        border: 1px solid rgba(197,160,89,0.2);
        box-shadow: inset 0 2px 4px rgba(0,0,0,0.2);
    }}
    
    .credit-bar {{
        height: 12px;
        background: linear-gradient(90deg, {theme['accent']}, {theme['accent2']}, {theme['accent']});
        background-size: 200% 100%;
        border-radius: 10px;
        transition: width 0.6s cubic-bezier(0.23, 1, 0.32, 1);
        box-shadow: 0 0 15px rgba(197,160,89,0.6), inset 0 1px 0 rgba(255,255,255,0.3);
        animation: gradientMove 3s linear infinite;
    }}
    
    @keyframes gradientMove {{
        0% {{ background-position: 0% 50%; }}
        100% {{ background-position: 200% 50%; }}
    }}

    /* === SECTION HEADER - PREMIUM === */
    .section-header {{
        color: {theme['accent']} !important;
        font-size: 14px;
        font-weight: 600;
        font-family: 'Inter', sans-serif;
        margin-top: 24px;
        margin-bottom: 12px;
        padding-bottom: 10px;
        border-bottom: 1px solid rgba(197,160,89,0.2);
        text-transform: uppercase;
        letter-spacing: 0.1em;
        display: flex;
        align-items: center;
        gap: 8px;
    }}
    
    .section-header::before {{
        content: '◆';
        font-size: 10px;
        color: {theme['accent']};
    }}

    /* === MOBILE NAV BUTTONS - PREMIUM === */
    .mobile-nav {{
        display: none;
        position: fixed;
        bottom: 24px;
        left: 50%;
        transform: translateX(-50%);
        z-index: 1000;
        background: linear-gradient(135deg, {theme['primary']} 0%, rgba(12,20,33,0.98) 100%);
        padding: 12px 24px;
        border-radius: 60px;
        box-shadow: var(--shadow-xl), 0 0 30px rgba(197,160,89,0.2);
        border: 1px solid rgba(197,160,89,0.3);
        backdrop-filter: blur(20px);
    }}
    
    .mobile-nav button {{
        background: linear-gradient(135deg, {theme['accent']} 0%, {theme['accent2']} 100%);
        color: {theme['primary']};
        border: none;
        padding: 12px 24px;
        margin: 0 6px;
        border-radius: 24px;
        font-size: 16px;
        cursor: pointer;
        font-weight: 600;
        transition: all var(--transition-fast);
        box-shadow: var(--shadow-sm);
    }}
    
    .mobile-nav button:hover {{
        transform: scale(1.05);
        box-shadow: var(--shadow-md);
    }}

    /* === PROGRESS BAR STYLING === */
    [data-testid='stProgress'] > div > div > div {{
        background: linear-gradient(90deg, {theme['accent']}, {theme['accent2']}, {theme['accent']}) !important;
        background-size: 200% 100% !important;
        animation: gradientMove 2s linear infinite !important;
        border-radius: 8px !important;
    }}
    
    [data-testid='stProgress'] {{
        background: rgba(0,0,0,0.2) !important;
        border-radius: 8px !important;
    }}

    /* === TOAST STYLING === */
    [data-testid='stToast'] {{
        background: linear-gradient(135deg, rgba(30,42,56,0.98) 0%, rgba(22,32,44,0.99) 100%) !important;
        border: 1px solid rgba(197,160,89,0.3) !important;
        border-radius: 12px !important;
        box-shadow: var(--shadow-lg), 0 0 30px rgba(197,160,89,0.15) !important;
        backdrop-filter: blur(20px) !important;
    }}
    
    [data-testid='stToast'] > div {{
        color: {text_primary} !important;
        font-family: 'Inter', sans-serif !important;
    }}

    /* === ALERT/WARNING BOXES === */
    [data-testid='stAlert'] {{
        background: linear-gradient(135deg, rgba(30,42,56,0.9) 0%, rgba(22,32,44,0.95) 100%) !important;
        border-radius: 12px !important;
        border-left: 4px solid !important;
        padding: 16px 20px !important;
    }}
    
    .stSuccess {{
        border-left-color: #10b981 !important;
    }}
    
    .stWarning {{
        border-left-color: #f59e0b !important;
    }}
    
    .stError {{
        border-left-color: #ef4444 !important;
    }}
    
    .stInfo {{
        border-left-color: {theme['accent']} !important;
    }}

    /* === DOWNLOAD BUTTON === */
    [data-testid='stDownloadButton'] > button {{
        background: linear-gradient(135deg, {theme['accent']} 0%, {theme['accent2']} 100%) !important;
        color: {theme['primary']} !important;
        font-weight: 700 !important;
        border: none !important;
        border-radius: 12px !important;
        padding: 16px 28px !important;
        font-size: 15px !important;
        transition: all var(--transition-normal) !important;
        box-shadow: var(--shadow-md) !important;
    }}
    
    [data-testid='stDownloadButton'] > button:hover {{
        transform: translateY(-3px) !important;
        box-shadow: var(--shadow-lg), var(--shadow-glow) !important;
    }}

    /* === LOADING SPINNER === */
    .loading-spinner {{
        display: inline-block;
        width: 20px;
        height: 20px;
        border: 2px solid rgba(197,160,89,0.3);
        border-radius: 50%;
        border-top-color: {theme['accent']};
        animation: spin 1s ease-in-out infinite;
    }}
    
    @keyframes spin {{
        to {{ transform: rotate(360deg); }}
    }}

    /* === TOOLTIP STYLES === */
    [data-tooltip] {{
        position: relative;
        cursor: help;
    }}
    
    [data-tooltip]::after {{
        content: attr(data-tooltip);
        position: absolute;
        bottom: 100%;
        left: 50%;
        transform: translateX(-50%) translateY(-8px);
        background: {theme['primary']};
        color: {theme['accent']};
        padding: 8px 12px;
        border-radius: 8px;
        font-size: 12px;
        white-space: nowrap;
        opacity: 0;
        pointer-events: none;
        transition: all var(--transition-fast);
        box-shadow: var(--shadow-lg);
        border: 1px solid rgba(197,160,89,0.3);
    }}
    
    [data-tooltip]:hover::after {{
        opacity: 1;
        transform: translateX(-50%) translateY(-4px);
    }}

    /* === BADGE / TAG STYLES === */
    .badge {{
        display: inline-flex;
        align-items: center;
        padding: 4px 12px;
        border-radius: 20px;
        font-size: 12px;
        font-weight: 600;
        font-family: 'Inter', sans-serif;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }}
    
    .badge-gold {{
        background: linear-gradient(135deg, {theme['accent']} 0%, {theme['accent2']} 100%);
        color: {theme['primary']};
    }}
    
    .badge-success {{
        background: linear-gradient(135deg, #10b981 0%, #059669 100%);
        color: white;
    }}

    /* === RESPONSIVE - ENHANCED === */
    @media (max-width: 768px) {{
        .main {{
            padding: 1rem 0.75rem;
        }}
        
        h1 {{
            font-size: 1.6rem !important;
        }}
        
        h1::after, h2::after {{
            width: 60px;
        }}
        
        .result-box {{
            padding: 22px !important;
            font-size: 15px !important;
            border-radius: 14px !important;
        }}
        
        .stButton>button {{
            font-size: 14px !important;
            padding: 14px 18px !important;
            border-radius: 10px !important;
        }}
        
        .login-card {{
            padding: 40px 25px;
            border-radius: 18px;
        }}
        
        .chat-user, .chat-ai {{
            padding: 14px 16px;
            border-radius: 14px 14px 4px 14px;
        }}
        
        .mobile-nav {{
            display: flex;
            align-items: center;
            gap: 8px;
        }}
        
        section[data-testid='stSidebar'] {{
            width: 280px !important;
        }}
    }}
    
    @media (max-width: 480px) {{
        .main {{
            padding: 0.75rem 0.5rem;
        }}
        
        h1 {{
            font-size: 1.4rem !important;
        }}
        
        .login-card {{
            padding: 30px 20px;
        }}
        
        .hero-title {{
            font-size: 1.8rem !important;
        }}
        
        .hero-subtitle {{
            font-size: 0.95rem !important;
        }}
    }}
    
    /* === PRINT STYLES === */
    @media print {{
        .stButton, .mobile-nav, section[data-testid='stSidebar'] {{
            display: none !important;
        }}
        
        .main {{
            background: white !important;
        }}
        
        .result-box {{
            border: 1px solid #ccc !important;
            box-shadow: none !important;
        }}
    }}
    </style>

    <script>
    // === KEYBOARD SHORTCUTS ===
    document.addEventListener('keydown', function(e) {{
        // Esc - Close modal
        if (e.key === 'Escape') {{
            const modal = document.getElementById('imageModal');
            if (modal) modal.style.display = 'none';
        }}
        
        // Ctrl+Enter - Trigger analysis
        if (e.ctrlKey && e.key === 'Enter') {{
            e.preventDefault();
            const analysisBtn = Array.from(document.querySelectorAll('.stButton button'))
                .find(btn => btn.textContent.includes('TAHLILNI BOSHLASH'));
            if (analysisBtn) analysisBtn.click();
        }}
        
        // Ctrl+S - Trigger download
        if (e.ctrlKey && e.key === 's') {{
            e.preventDefault();
            const downloadBtn = document.querySelector('[data-testid="stDownloadButton"]');
            if (downloadBtn) downloadBtn.click();
        }}
    }});

    // === IMAGE MODAL ===
    function setupModal() {{
        const images = document.querySelectorAll('.magnifier-container img');
        const modal = document.getElementById('imageModal');
        const modalImg = document.getElementById('modalImage');
        const closeBtn = document.querySelector('.modal-close');
        
        images.forEach(img => {{
            img.onclick = function() {{
                if (modal && modalImg) {{
                    modal.style.display = "block";
                    modalImg.src = this.src;
                }}
            }}
        }});
        
        if (closeBtn) {{
            closeBtn.onclick = function() {{
                modal.style.display = "none";
            }}
        }}
        
        if (modal) {{
            modal.onclick = function(e) {{
                if (e.target === modal) {{
                    modal.style.display = "none";
                }}
            }}
        }}
    }}

    // === MOBILE SWIPE GESTURES ===
    let touchstartX = 0;
    let touchendX = 0;
    
    document.addEventListener('touchstart', e => {{
        touchstartX = e.changedTouches[0].screenX;
    }});
    
    document.addEventListener('touchend', e => {{
        touchendX = e.changedTouches[0].screenX;
        handleSwipe();
    }});
    
    function handleSwipe() {{
        const threshold = 100;
        if (touchendX < touchstartX - threshold) {{
            // Swipe left - next page
            const nextBtn = document.getElementById('mobileNext');
            if (nextBtn) nextBtn.click();
        }}
        if (touchendX > touchstartX + threshold) {{
            // Swipe right - previous page
            const prevBtn = document.getElementById('mobilePrev');
            if (prevBtn) prevBtn.click();
        }}
    }}

    // Run after Streamlit renders
    setTimeout(setupModal, 1000);
    setInterval(setupModal, 2000); // Refresh for dynamic content
    </script>

    <!-- Modal HTML -->
    <div id="imageModal" class="modal">
        <span class="modal-close">&times;</span>
        <img class="modal-content" id="modalImage">
    </div>
""", unsafe_allow_html=True)

# --- 2. CORE SERVICES (SUPABASE & AI) ---
try:
    db = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])
    CORRECT_PASSWORD = st.secrets["APP_PASSWORD"]
    GEMINI_KEY = st.secrets["GEMINI_API_KEY"]
except:
    st.error("Secrets sozlanmagan!")
    st.stop()

# ==========================================
# DEMO MODE (PITCH / TANLOV UCHUN)
# ==========================================
# Eslatma:
# - DEMO_MODE=True bo'lsa, tizimga kirishda parol so'ralmaydi (faqat email).
# - Yangi foydalanuvchi birinchi kirishda avtomatik 50 kredit bilan yaratiladi.
DEMO_MODE = True
DEFAULT_DEMO_CREDITS = 50

def ensure_demo_user(email: str) -> None:
    """Agar foydalanuvchi mavjud bo'lmasa, uni demo krediti bilan yaratadi.

    Muhim:
    - Faqat tanlov/pitch jarayonida qulaylik uchun.
    - Jadval sxemasi noma'lum bo'lgani uchun minimal fieldlar bilan insert qilinadi.
    """
    try:
        # Avval mavjudligini tekshiramiz
        res = db.table("profiles").select("email").eq("email", email).execute()
        if not res.data:
            db.table("profiles").insert({
                "email": email,
                "credits": DEFAULT_DEMO_CREDITS,
            }).execute()
    except Exception:
        # Demo rejimda login to'xtab qolmasligi uchun xatoni yumshoq usulda yutamiz.
        # Kreditlar baribir 0 bo'lib qolishi mumkin (bazaga yozilmasa).
        pass

if not st.session_state.auth:
    # === ENHANCED LOGIN PAGE ===
    st.markdown("<br>", unsafe_allow_html=True)
    _, col_mid, _ = st.columns([1, 2, 1])
    with col_mid:
        st.markdown(f"""
            <div class='login-card'>
                <div class='hero-title'>🏛 Manuscript AI</div>
                <div class='hero-subtitle'>
                    Qadimiy qo'lyozmalarni raqamli tahlil qilish va transliteratsiya 
                    qilish uchun sun'iy intellekt asosidagi platforma
                </div>
            </div>
        """, unsafe_allow_html=True)
        
        st.markdown("<h2 style='margin-top:30px;'>🔐 Tizimga Kirish</h2>", unsafe_allow_html=True)
        
        email_in = st.text_input("📧 Email manzili", placeholder="example@domain.com")

        # DEMO_MODE=True bo'lsa, parol maydoni ko'rsatilmaydi (pitch/tanlov uchun qulay).
        if not DEMO_MODE:
            pwd_in = st.text_input("🔑 Parol", type="password", placeholder="Parolingizni kiriting")

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button("✨ TIZIMGA KIRISH"):
            if not email_in:
                st.warning("⚠️ Iltimos, email manzilini kiriting")
            elif DEMO_MODE:
                st.session_state.auth = True
                st.session_state.u_email = email_in.strip().lower()
                ensure_demo_user(st.session_state.u_email)
                st.toast("✅ Tizimga muvaffaqiyatli kirdingiz!", icon="🎉")
                st.rerun()
            else:
                if pwd_in == CORRECT_PASSWORD:
                    st.session_state.auth = True
                    st.session_state.u_email = email_in.strip().lower()
                    st.toast("✅ Muvaffaqiyatli kirdingiz!", icon="🎉")
                    st.rerun()
                else:
                    st.error("❌ Parol noto'g'ri! Iltimos, qaytadan urinib ko'ring.")
    st.stop()

# --- AI ENGINE (UNCHANGED - DO NOT MODIFY) ---
genai.configure(api_key=GEMINI_KEY)
system_instruction = "Siz Manuscript AI mutaxassisiz. Tadqiqotchi d87809889-dot muallifligida ilmiy tahlillar qilasiz."
safety_settings = {
    HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_HATE_SPEECH: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_SEXUALLY_EXPLICIT: HarmBlockThreshold.BLOCK_NONE,
    HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE
}

# Generation config - aniqlik uchun optimallashtirilgan
generation_config = genai.GenerationConfig(
    temperature=0.1,  # 0.1 = deterministik lekin qotib qolmaydi, nozik tafovutlarni sezadi
    top_p=0.95,
    top_k=40,
    max_output_tokens=8192,  # 4096→8192: barcha 7 bo'lim to'liq chiqadi
)

model = genai.GenerativeModel(
    model_name="gemini-flash-latest",
    system_instruction=system_instruction,
    safety_settings=safety_settings,
    generation_config=generation_config
)

# ==========================================
# 3. YORDAMCHI FUNKSIYALAR (UNCHANGED)
# ==========================================

# ==========================================
# 3.1 ADVANCED IMAGE PREPROCESSING
# ==========================================
def safe_denoise(img: Image.Image) -> Image.Image:
    """Xavfsiz shovqin olib tashlash - PIL bilan"""
    try:
        # Median filter - shovqinni kamaytiradi, chetlarni saqlaydi
        return img.filter(ImageFilter.MedianFilter(size=3))
    except Exception:
        return img  # Xato bo'lsa original qaytaradi

def safe_deskew(img: Image.Image) -> Image.Image:
    """Xavfsiz qiyshiqlikni to'g'rilash - oddiy usul"""
    try:
        # Oddiy autorotate - EXIF asosida
        return ImageOps.exif_transpose(img) or img
    except Exception:
        return img

def adaptive_binarize(img: Image.Image) -> Image.Image:
    """Matn/fon ajratish - adaptive threshold"""
    try:
        # Grayscale bo'lishi kerak
        if img.mode != 'L':
            img = img.convert('L')
        
        # Oddiy threshold
        threshold = 128
        return img.point(lambda p: 255 if p > threshold else 0)
    except Exception:
        return img

def optimal_resize(img: Image.Image, target_size: int = 1800) -> Image.Image:
    """Optimal o'lchamga keltirish - SIFAT uchun optimallashtirilgan"""
    try:
        w, h = img.size
        max_dim = max(w, h)
        
        # Agar kichik bo'lsa - kattalashtiramiz (ko'proq detal uchun)
        if max_dim < 900:
            scale = 1200 / max_dim  # 1000→1200: kichik rasmlarni aniqroq ko'rsatadi
            new_w, new_h = int(w * scale), int(h * scale)
            return img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        
        # Agar katta bo'lsa - kichraytramiz (lekin sifatni saqlaymiz)
        if max_dim > 2200:  # 1600→2200: katta rasmlarni kamroq kichraytiramiz
            scale = target_size / max_dim
            new_w, new_h = int(w * scale), int(h * scale)
            return img.resize((new_w, new_h), Image.Resampling.LANCZOS)
        
        return img
    except Exception:
        return img

def enhance_image_for_ai(img: Image.Image) -> Image.Image:
    """Rasmni AI tahlili uchun optimallashtirish
    - RANGLI saqlanadi (siyoh rangi muhim)
    - Engil kontrast/keskinlik (o'chgan siyoh uchun)
    - Grayscale QILINMAYDI
    """
    try:
        # 1. Optimal o'lchamga keltirish
        img = optimal_resize(img, target_size=1800)
        
        # 2. EXIF rotatsiyasini to'g'rilash
        img = ImageOps.exif_transpose(img) or img
        
        # 3. Engil kontrast (rangli rasmda ham ishlaydi, o'chgan siyohni aniqlashtiradi)
        img = ImageEnhance.Contrast(img).enhance(1.2)
        
        # 4. Engil keskinlik (harflar chegarasini biroz aniqlashtiradi)
        img = ImageEnhance.Sharpness(img).enhance(1.15)
        
        return img
    except Exception:
        return img

# ==========================================
# 3.2 SMART CROPPING - Katta rasmlar uchun
# ==========================================
def should_split_image(img: Image.Image) -> bool:
    """Rasm bo'linishi kerakmi? - FAQAT juda katta rasmlar uchun"""
    w, h = img.size
    # Chegarani oshirdik: oddiy rasmlarni BO'LMAYMIZ (qatorlar aralashmasligi uchun)
    # Faqat juda katta rasmlarni bo'lamiz (5000+ piksel)
    return h > 5000 or w > 5000

def split_image_smart(img: Image.Image) -> list:
    """Rasmni aqlli bo'laklarga bo'lish - OVERLAP KAMAYTIRILDI"""
    try:
        w, h = img.size
        crops = []
        
        # Vertikal bo'lish (asosan)
        if h > w * 1.5:  # Vertikal rasm
            # 2 qismga bo'lish (FAQAT 3% overlap - qatorlar takrorlanmasligi uchun)
            mid = h // 2
            overlap = int(h * 0.03)  # 10%→3%: minimal xavfsizlik chegarasi
            
            crop1 = img.crop((0, 0, w, mid + overlap))
            crop2 = img.crop((0, mid - overlap, w, h))
            crops = [crop1, crop2]
        
        # Gorizontal bo'lish
        elif w > h * 1.5:  # Gorizontal rasm
            mid = w // 2
            overlap = int(w * 0.03)  # 10%→3%
            
            crop1 = img.crop((0, 0, mid + overlap, h))
            crop2 = img.crop((mid - overlap, 0, w, h))
            crops = [crop1, crop2]
        
        # Kvadrat - 4 qismga bo'lmaslik, FAQAT 2 qismga (aniqroq)
        else:
            mid_h = h // 2
            overlap = int(h * 0.03)
            
            crop1 = img.crop((0, 0, w, mid_h + overlap))
            crop2 = img.crop((0, mid_h - overlap, w, h))
            crops = [crop1, crop2]  # 4→2: kamroq bo'lak = kamroq xato
        
        return crops if crops else [img]
    except Exception:
        return [img]

def merge_results(results: list) -> str:
    """Bo'lak natijalarini birlashtirish va yaxlit ko'rinishga keltirish"""
    if not results:
        return ""
    if len(results) == 1:
        return results[0]
    
    # Bo'laklarni birlashtirish
    all_transliterations = []
    all_translations = []
    all_notes = []
    
    for r in results:
        # Har bir qismdan bo'limlarni ajratib olish
        text = r.strip()
        
        # Transliteratsiya qismini topish
        trans_start = text.lower().find("transliter")
        tarj_start = text.lower().find("tarjima") if text.lower().find("tarjima") > 0 else text.lower().find("translation")
        izoh_start = text.lower().find("izoh") if text.lower().find("izoh") > 0 else text.lower().find("note")
        
        if trans_start >= 0 and tarj_start > trans_start:
            all_transliterations.append(text[trans_start:tarj_start].strip())
        
        if tarj_start >= 0 and izoh_start > tarj_start:
            all_translations.append(text[tarj_start:izoh_start].strip())
        elif tarj_start >= 0:
            all_translations.append(text[tarj_start:].strip())
        
        if izoh_start >= 0:
            all_notes.append(text[izoh_start:].strip())
    
    # Yaxlit natija shakllantirish
    merged = "📜 **YAXLIT TAHLIL NATIJASI**\n\n"
    merged += "---\n\n"
    
    if all_transliterations:
        merged += "## 1. TRANSLITERATSIYA (Asl yozuv)\n\n"
        for i, t in enumerate(all_transliterations, 1):
            # Sarlavhalarni olib tashlash
            clean = t.replace("**TRANSLITERATSIYA**:", "").replace("**TRANSLITERATION**:", "").replace("1.", "").strip()
            if clean:
                merged += f"{clean}\n\n"
        merged += "---\n\n"
    
    if all_translations:
        merged += "## 2. TO'LIQ TARJIMA (O'zbek tilida)\n\n"
        for i, t in enumerate(all_translations, 1):
            clean = t.replace("**TARJIMA**:", "").replace("**TRANSLATION**:", "").replace("**TO'LIQ TARJIMA**:", "").replace("2.", "").strip()
            if clean:
                merged += f"{clean}\n\n"
        merged += "---\n\n"
    
    if all_notes:
        merged += "## 3. IZOHLAR\n\n"
        for i, t in enumerate(all_notes, 1):
            clean = t.replace("**IZOHLAR**:", "").replace("**NOTES**:", "").replace("3.", "").strip()
            if clean:
                merged += f"{clean}\n\n"
    
    return merged

# ==========================================
# 3.3 QUALITY-BASED RETRY
# ==========================================
def assess_quality(response_text: str) -> dict:
    """Javob sifatini baholash - KENGAYTIRILGAN VERSIYA"""
    if not response_text:
        return {"score": 0, "reason": "Bo'sh javob", "retry": True, "details": {}}
    
    text = response_text.lower()
    
    # Skor hisoblash
    score = 100
    reasons = []
    details = {}
    
    # 1. Noaniqlik belgilari tekshirish (ANIQROQ CHEGARALAR)
    unclear_count = text.count("[?]") + text.count("unclear") + text.count("noaniq") + text.count("[...]")
    details['unclear_marks'] = unclear_count
    
    if unclear_count > 20:
        score -= 40
        reasons.append(f"{unclear_count} noaniq belgi (juda ko'p - rasmni yaxshilang)")
    elif unclear_count > 12:
        score -= 25
        reasons.append(f"{unclear_count} noaniq belgi (ko'p)")
    elif unclear_count > 6:
        score -= 12
        reasons.append(f"{unclear_count} noaniq belgi (o'rtacha)")
    elif unclear_count > 0:
        score -= 3
        reasons.append(f"{unclear_count} noaniq belgi (oz)")
    
    # 2. Javob uzunligi tekshirish (ANIQROQ)
    word_count = len(response_text.split())
    details['word_count'] = word_count
    
    if word_count < 30:
        score -= 40
        reasons.append("Juda qisqa javob (kam ma'lumot)")
    elif word_count < 80:
        score -= 20
        reasons.append("Qisqa javob (to'liqroq bo'lishi kerak)")
    elif word_count < 150:
        score -= 5
        reasons.append("Qisqaroq javob")
    else:
        score += 5  # Bonus batafsil javob uchun
    
    # 3. MAJBURIY bo'limlar tekshirish
    required_sections = [
        ("transliteratsiya", "Transliteratsiya bo'limi yo'q", 20),
        ("tarjima", "Tarjima bo'limi yo'q", 20),
        ("leksik", "Leksik tahlil yo'q", 15),
        ("identifikatsiya", "Manba identifikatsiyasi yo'q", 10),
        ("izoh", "Izohlar bo'limi yo'q", 10),
    ]
    
    missing_sections = []
    for keyword, error_msg, penalty in required_sections:
        if keyword not in text:
            score -= penalty
            reasons.append(error_msg)
            missing_sections.append(keyword)
    details['missing_sections'] = missing_sections
    
    # 4. Jadval formati tekshirish (ANIQROQ)
    table_rows = response_text.count("|")
    details['table_rows'] = table_rows
    
    if table_rows < 3:
        score -= 15
        reasons.append("Jadval yo'q yoki to'liq emas")
    elif table_rows < 8:
        score -= 5
        reasons.append("Jadval qisqa (kamida 5-10 qator bo'lishi kerak)")
    
    # 5. Xato xabarlari tekshirish
    error_keywords = ["error", "xato", "imkonsiz", "o'qib bo'lmaydi", "ko'rinmaydi", "butunlay", "yo'q"]
    found_errors = []
    for kw in error_keywords:
        if kw in text and text.count(kw) > 2:  # 2 martadan ko'p
            found_errors.append(kw)
    
    if found_errors:
        score -= 15
        reasons.append(f"Ko'p xato belgilari: {', '.join(found_errors[:3])}")
    details['error_keywords'] = found_errors
    
    # 6. Aniqlik bahosi tekshirish (YANGI: sonlarni tekshirish)
    percent_count = response_text.count("%")
    details['percent_marks'] = percent_count
    
    if percent_count == 0:
        score -= 10
        reasons.append("Aniqlik foizlari yo'q")
    elif percent_count < 3:
        score -= 5
        reasons.append("Kamida 3 ta aniqlik foizi bo'lishi kerak")
    
    # 7. YANGI: Bo'lim sarlavhalari mavjudligi
    section_headers = response_text.count("##")
    details['section_headers'] = section_headers
    
    if section_headers < 5:
        score -= 10
        reasons.append(f"Faqat {section_headers} ta bo'lim sarlavhasi (kamida 5-7 ta kerak)")
    
    # 8. YANGI: Javob strukturasini tekshirish
    has_good_structure = ("##" in response_text and 
                         response_text.count("\n\n") > 5 and
                         len(response_text.split("\n")) > 20)
    if not has_good_structure:
        score -= 8
        reasons.append("Yomon formatlangan javob")
    details['has_structure'] = has_good_structure
    
    # Final skorni hisoblash
    final_score = max(0, min(100, score))
    
    # Sifat darajasini aniqlash
    if final_score >= 85:
        quality_level = "A'lo"
    elif final_score >= 70:
        quality_level = "Yaxshi"
    elif final_score >= 55:
        quality_level = "Qoniqarli"
    else:
        quality_level = "Past"
    
    return {
        "score": final_score,
        "level": quality_level,
        "reason": ", ".join(reasons) if reasons else "Mukammal sifat",
        "retry": final_score < 55,  # 55 dan past bo'lsa qayta urinish
        "details": details
    }

def generate_quality_report(quality: dict, theme: dict) -> str:
    """Sifat hisobotini HTML formatida yaratish"""
    score = quality['score']
    level = quality.get('level', 'Noma\'lum')
    details = quality.get('details', {})
    
    # Rang va emoji tanlash
    if score >= 85:
        emoji = "🏆"
        color = "#10b981"
        bg_color = "rgba(16, 185, 129, 0.1)"
    elif score >= 70:
        emoji = "✅"
        color = "#3b82f6"
        bg_color = "rgba(59, 130, 246, 0.1)"
    elif score >= 55:
        emoji = "⚠️"
        color = "#f59e0b"
        bg_color = "rgba(245, 158, 11, 0.1)"
    else:
        emoji = "❌"
        color = "#ef4444"
        bg_color = "rgba(239, 68, 68, 0.1)"
    
    # Batafsil metrikalar
    metrics_html = ""
    if details:
        metrics_html = "<div style='margin-top:12px; padding:12px; background:rgba(0,0,0,0.2); border-radius:8px;'>"
        metrics_html += "<p style='font-size:12px; color:rgba(255,255,255,0.7); margin:4px 0;'><b>📊 Metrikalar:</b></p>"
        
        if 'word_count' in details:
            metrics_html += f"<p style='font-size:11px; margin:2px 0;'>• So'zlar soni: {details['word_count']}</p>"
        if 'unclear_marks' in details:
            metrics_html += f"<p style='font-size:11px; margin:2px 0;'>• Noaniq belgilar: {details['unclear_marks']}</p>"
        if 'table_rows' in details:
            metrics_html += f"<p style='font-size:11px; margin:2px 0;'>• Jadval qatorlari: {details['table_rows']}</p>"
        if 'section_headers' in details:
            metrics_html += f"<p style='font-size:11px; margin:2px 0;'>• Bo'limlar soni: {details['section_headers']}</p>"
        if 'missing_sections' in details and details['missing_sections']:
            metrics_html += f"<p style='font-size:11px; margin:2px 0; color:#ef4444;'>• Yo'q bo'limlar: {', '.join(details['missing_sections'][:3])}</p>"
        
        metrics_html += "</div>"
    
    return f"""
<div style='background:{bg_color}; padding:16px; border-radius:12px; 
            border-left:4px solid {color}; margin:10px 0; backdrop-filter:blur(10px);'>
    <div style='display:flex; justify-content:space-between; align-items:center;'>
        <h4 style='margin:0; color:{color}; font-size:18px;'>{emoji} {level} Sifat</h4>
        <span style='background:{color}; color:white; padding:6px 12px; border-radius:20px; 
                     font-weight:700; font-size:16px;'>{score}%</span>
    </div>
    <p style='margin:8px 0 0 0; font-size:14px; color:rgba(255,255,255,0.8);'>{quality['reason']}</p>
    {metrics_html}
</div>
"""

def analyze_with_retry(model, prompt: str, img: Image.Image, max_retries: int = 2) -> tuple:
    """1 ta sifatli so'rov + faqat kerak bo'lsa qayta urinish
    
    OLDIN: har doim 2-3 ta API chaqiruv (dual_pass + retry)
    ENDI:  1 ta chaqiruv, faqat sifat < 50 bo'lsa 1 marta qayta urinish
    NATIJA: 2-3x kamroq API sarfi, tezlik 2x oshadi
    """
    
    # === BIRINCHI SO'ROV: rangli original rasm ===
    result = None
    quality = {"score": 0, "reason": "Natija olinmadi", "retry": True}
    
    try:
        processed_img = enhance_image_for_ai(img)
        payload = img_to_png_payload(processed_img)
        resp = model.generate_content([prompt, payload])
        
        if resp.candidates and resp.candidates[0].content.parts:
            result = resp.text
            quality = assess_quality(result)
    except Exception:
        pass
    
    # Agar yaxshi natija bo'lsa — darhol qaytarish (1 ta API chaqiruv)
    if result and quality["score"] >= 50:
        result = post_process_result(result)
        return (result, quality, 1)
    
    # === QAYTA URINISH: faqat sifat < 50 bo'lsa (grayscale + kontrast) ===
    try:
        retry_img = ImageOps.grayscale(img)
        retry_img = optimal_resize(retry_img, target_size=1800)
        retry_img = ImageOps.autocontrast(retry_img, cutoff=1)
        retry_img = ImageEnhance.Contrast(retry_img).enhance(1.8)
        
        payload = img_to_png_payload(retry_img)
        resp = model.generate_content([prompt, payload])
        
        if resp.candidates and resp.candidates[0].content.parts:
            retry_result = resp.text
            retry_quality = assess_quality(retry_result)
            
            # Yangi natija yaxshiroq bo'lsa — uni olish
            if retry_quality["score"] > quality["score"]:
                return (post_process_result(retry_result), retry_quality, 2)
    except Exception:
        pass
    
    # Original natijani qaytarish
    if result:
        result = post_process_result(result)
    return (result, quality, 1)

def img_to_png_payload(img: Image.Image):
    """Rasmni PNG formatda AIga yuborish (lossless - nuqtalar saqlanadi)"""
    buffered = io.BytesIO()
    if img.mode == 'RGBA':
        img = img.convert('RGB')
    img.save(buffered, format="PNG", optimize=False, compress_level=3)
    return {"mime_type": "image/png", "data": base64.b64encode(buffered.getvalue()).decode("utf-8")}


def post_process_result(result_text: str) -> str:
    """AI natijasini tozalash va formatlash - POST-PROCESSING"""
    if not result_text:
        return result_text
    
    text = result_text.strip()
    
    # 1. Bo'sh qatorlarni tozalash (3+ bo'sh qator → 2 ta)
    import re
    text = re.sub(r'\n{4,}', '\n\n\n', text)
    
    # 2. Jadval formatini to'g'rilash
    lines = text.split('\n')
    fixed_lines = []
    
    for line in lines:
        # Jadval qatorlarini tekshirish
        if '|' in line:
            # Boshi va oxiridagi | ni tekshirish
            stripped = line.strip()
            if stripped and not stripped.startswith('|'):
                stripped = '| ' + stripped
            if stripped and not stripped.endswith('|'):
                stripped = stripped + ' |'
            fixed_lines.append(stripped)
        else:
            fixed_lines.append(line)
    
    text = '\n'.join(fixed_lines)
    
    # 3. Bo'lim sarlavhalarini standartlashtirish
    section_fixes = [
        (r'##\s*1[\.\)]\s*', '## 1. '),
        (r'##\s*2[\.\)]\s*', '## 2. '),
        (r'##\s*3[\.\)]\s*', '## 3. '),
        (r'##\s*4[\.\)]\s*', '## 4. '),
        (r'##\s*5[\.\)]\s*', '## 5. '),
        (r'##\s*6[\.\)]\s*', '## 6. '),
        (r'##\s*7[\.\)]\s*', '## 7. '),
    ]
    
    for pattern, replacement in section_fixes:
        text = re.sub(pattern, replacement, text)
    
    # 4. Noaniq belgilarni standartlashtirish
    text = text.replace('[unclear]', '[?]')
    text = text.replace('[noaniq]', '[?]')
    text = text.replace('[ ? ]', '[?]')
    text = text.replace('[  ?  ]', '[?]')
    
    # 5. Foiz belgilarini tozalash
    text = re.sub(r'(\d+)\s*%', r'\1%', text)
    
    return text


# ==========================================
# DUAL-PASS ANALYZE (ARCHIVED - NO LONGER USED)
# ==========================================
# Oldingi versiyada 2-3 marta API chaqirilardi.
# Hozir single-pass + conditional retry ishlatiladi (kamroq API chaqiruvi).
# Bu funksiya faqat arxiv uchun saqlanmoqda.

# def dual_pass_analyze(model, prompt: str, img: Image.Image) -> tuple:
#     """Dual-Pass Tahlil - 2 xil usulda tahlil qilib, eng yaxshisini tanlash"""
#     results = []
#     
#     # === PASS 1: Standart preprocessing ===
#     try:
#         processed_img1 = enhance_image_for_ai(img)
#         payload1 = img_to_png_payload(processed_img1)
#         resp1 = model.generate_content([prompt, payload1])
#         
#         if resp1.candidates and resp1.candidates[0].content.parts:
#             result1 = resp1.text
#             quality1 = assess_quality(result1)
#             results.append({
#                 'text': result1,
#                 'quality': quality1,
#                 'method': 'standard'
#             })
#     except Exception as e:
#         pass  # Xato bo'lsa keyingi usulga o'tamiz
#     
#     # === PASS 2: Yuqori kontrast preprocessing ===
#     try:
#         processed_img2 = enhance_image_for_ai(img)
#         # Qo'shimcha kontrast va keskinlik
#         processed_img2 = ImageEnhance.Contrast(processed_img2).enhance(1.5)
#         processed_img2 = ImageEnhance.Sharpness(processed_img2).enhance(1.3)
#         
#         payload2 = img_to_png_payload(processed_img2)
#         resp2 = model.generate_content([prompt, payload2])
#         
#         if resp2.candidates and resp2.candidates[0].content.parts:
#             result2 = resp2.text
#             quality2 = assess_quality(result2)
#             results.append({
#                 'text': result2,
#                 'quality': quality2,
#                 'method': 'high_contrast'
#             })
#     except Exception as e:
#         pass
#     
#     # === ENG YAXSHI NATIJANI TANLASH ===
#     if not results:
#         return (None, {"score": 0, "reason": "Hech qanday natija olinmadi", "retry": False}, 0)
#     
#     # Eng yuqori sifatli natijani tanlash
#     best = max(results, key=lambda x: x['quality']['score'])
#     
#     # Post-processing qo'llash
#     final_text = post_process_result(best['text'])
#     
#     # Agar ikkala natija ham yaxshi bo'lsa, uzunroq natijani olish
#     if len(results) == 2:
#         score_diff = abs(results[0]['quality']['score'] - results[1]['quality']['score'])
#         if score_diff < 10:  # Skorlar yaqin bo'lsa
#             # Uzunroq javobni olish (ko'proq ma'lumot)
#             longer = max(results, key=lambda x: len(x['text']))
#             final_text = post_process_result(longer['text'])
#             best = longer
#     
#     return (final_text, best['quality'], len(results))

def fetch_live_credits(email: str):
    try:
        res = db.table("profiles").select("credits").eq("email", email).single().execute()
        return res.data["credits"] if res.data else 0
    except:
        return 0

def use_credit_atomic(email: str, count: int = 1):
    curr = fetch_live_credits(email)
    if curr >= count:
        db.table("profiles").update({"credits": curr - count}).eq("email", email).execute()
        return True
    return False

@st.cache_data(show_spinner=False)
def render_page(file_content, page_idx, scale, is_pdf):
    try:
        if is_pdf:
            pdf = pdfium.PdfDocument(file_content)
            img = pdf[page_idx].render(scale=scale).to_pil()
            pdf.close()
            return img
        return Image.open(io.BytesIO(file_content))
    except:
        return None

# ==========================================
# UI CONSTANTS (DEMO METRICS - NOT LIVE DATA)
# ==========================================
# NOTE: These are placeholder values for demo purposes
DEMO_MANUSCRIPTS_ANALYZED = 2840
DEMO_LANGUAGES_SUPPORTED = 12
DEMO_AVG_TIME_MINUTES = 2.5
DEMO_ACCURACY_RATE = 92.1
DEMO_ACTIVE_USERS = 180
DEMO_COUNTRIES = 4

# ==========================================
# LANDING PAGE FUNCTION (UI ONLY)
# ==========================================
def render_landing_page():
    """Render professional landing page"""
    
    # Hero Section
    st.markdown(f"""
        <div style='text-align:center; padding:70px 30px 50px; 
                    background: linear-gradient(145deg, rgba(12,20,33,0.98) 0%, rgba(22,32,44,0.99) 100%);
                    border-radius:24px; 
                    box-shadow: 0 20px 60px rgba(0,0,0,0.3), 0 0 80px rgba(197,160,89,0.06);
                    border: 1px solid rgba(197,160,89,0.15);
                    margin-bottom:40px;
                    position: relative;
                    overflow: hidden;'>
            <div style='position:absolute; top:0; left:0; right:0; height:3px;
                        background: linear-gradient(90deg, {theme['accent']}, {theme['accent2']}, {theme['accent']});
                        background-size: 200% 100%;
                        animation: shimmer 3s ease-in-out infinite;'></div>
            <div style='font-size:64px; margin-bottom:20px; animation: float 3s ease-in-out infinite;'>🏛️</div>
            <h1 style='font-size:clamp(2.5rem, 6vw, 3.8rem); margin-bottom:16px; border:none;
                       background: linear-gradient(135deg, {theme['accent']}, {theme['accent2']});
                       -webkit-background-clip: text; -webkit-text-fill-color: transparent;
                       background-clip: text; font-family: Playfair Display, serif;'>
                Manuscript AI
            </h1>
            <p style='font-size:clamp(1.1rem, 2.5vw, 1.4rem); color: rgba(253,250,241,0.6);
                      margin-bottom:40px; line-height:1.7; max-width:600px; margin-left:auto; margin-right:auto;'>
                Qadimiy qo'lyozmalarni raqamli tahlil qilish va transliteratsiya qilish uchun<br>
                sun'iy intellekt asosidagi platforma
            </p>
            <div style='display:flex; justify-content:center; gap:30px; flex-wrap:wrap; margin-bottom:10px;'>
                <div style='text-align:center;'>
                    <div style='font-size:2rem; font-weight:700; color:{theme['accent']};
                                font-family: Playfair Display, serif;'>{DEMO_MANUSCRIPTS_ANALYZED}+</div>
                    <div style='font-size:0.8rem; color:rgba(253,250,241,0.45); text-transform:uppercase;
                                letter-spacing:0.1em; margin-top:4px;'>Tahlil qilingan</div>
                </div>
                <div style='width:1px; background:rgba(197,160,89,0.2); align-self:stretch;'></div>
                <div style='text-align:center;'>
                    <div style='font-size:2rem; font-weight:700; color:{theme['accent']};
                                font-family: Playfair Display, serif;'>{DEMO_ACCURACY_RATE}%</div>
                    <div style='font-size:0.8rem; color:rgba(253,250,241,0.45); text-transform:uppercase;
                                letter-spacing:0.1em; margin-top:4px;'>Aniqlik</div>
                </div>
                <div style='width:1px; background:rgba(197,160,89,0.2); align-self:stretch;'></div>
                <div style='text-align:center;'>
                    <div style='font-size:2rem; font-weight:700; color:{theme['accent']};
                                font-family: Playfair Display, serif;'>{DEMO_LANGUAGES_SUPPORTED}</div>
                    <div style='font-size:0.8rem; color:rgba(253,250,241,0.45); text-transform:uppercase;
                                letter-spacing:0.1em; margin-top:4px;'>Tillar</div>
                </div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    # Feature cards
    f1, f2, f3 = st.columns(3)
    feat_css = f"""background: linear-gradient(145deg, rgba(30,42,56,0.7) 0%, rgba(22,32,44,0.9) 100%);
                   border: 1px solid rgba(197,160,89,0.15); border-radius:16px; padding:28px 22px;
                   text-align:center; box-shadow: 0 8px 32px rgba(0,0,0,0.15); min-height:180px;"""
    with f1:
        st.markdown(f"""<div style='{feat_css}'>
            <div style='font-size:36px; margin-bottom:12px;'>🔍</div>
            <h4 style='color:{theme['accent']}; margin:0 0 8px 0; font-size:1.1rem; border:none; padding:0;'>AI Tahlil</h4>
            <p style='color:rgba(253,250,241,0.5); font-size:0.85rem; margin:0; line-height:1.5;'>
                Gemini Vision orqali qo'lyozmalarni avtomatik tahlil qilish
            </p>
        </div>""", unsafe_allow_html=True)
    with f2:
        st.markdown(f"""<div style='{feat_css}'>
            <div style='font-size:36px; margin-bottom:12px;'>🔤</div>
            <h4 style='color:{theme['accent']}; margin:0 0 8px 0; font-size:1.1rem; border:none; padding:0;'>Transliteratsiya</h4>
            <p style='color:rgba(253,250,241,0.5); font-size:0.85rem; margin:0; line-height:1.5;'>
                Arab-Fors yozuvidan lotin alifbosiga qator-qator o'tkazish
            </p>
        </div>""", unsafe_allow_html=True)
    with f3:
        st.markdown(f"""<div style='{feat_css}'>
            <div style='font-size:36px; margin-bottom:12px;'>📚</div>
            <h4 style='color:{theme['accent']}; margin:0 0 8px 0; font-size:1.1rem; border:none; padding:0;'>Akademik Tahlil</h4>
            <p style='color:rgba(253,250,241,0.5); font-size:0.85rem; margin:0; line-height:1.5;'>
                Leksik-semantik tahlil va akademik izohlar
            </p>
        </div>""", unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # CTA Button
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("🚀 BOSHLASH", key="start_btn", use_container_width=True):
            st.session_state.show_landing = False
            st.rerun()
    
    st.markdown(f"""
        <div style='text-align:center; margin-top:50px; padding:20px;
                    border-top: 1px solid rgba(197,160,89,0.1);'>
            <p style='color:rgba(253,250,241,0.3); font-size:0.8rem; margin:0;'>
                Tadqiqot: d87809889-dot · 📧 {st.session_state.u_email}
            </p>
        </div>
    """, unsafe_allow_html=True)

# ==========================================
# 4. TADQIQOT INTERFEYSI
# ==========================================
with st.sidebar:
    # === SIDEBAR HEADER ===
    st.markdown(f"""
        <div style='text-align:center; padding:16px 0 12px; margin-bottom:8px;'>
            <div style='font-size:32px; margin-bottom:6px;'>🏛️</div>
            <div style='font-family: Playfair Display, serif; font-size:1.4rem; font-weight:700;
                        background: linear-gradient(135deg, {theme['accent']}, {theme['accent2']});
                        -webkit-background-clip: text; -webkit-text-fill-color: transparent;
                        background-clip: text;'>Manuscript AI</div>
        </div>
    """, unsafe_allow_html=True)
    
    # User info compact
    st.markdown(f"""
        <div style='background: rgba(197,160,89,0.08); border-radius:10px; padding:10px 14px;
                    border: 1px solid rgba(197,160,89,0.12); margin-bottom:12px;
                    display:flex; align-items:center; gap:10px;'>
            <div style='width:32px; height:32px; border-radius:50%;
                        background: linear-gradient(135deg, {theme['accent']}, {theme['accent2']});
                        display:flex; align-items:center; justify-content:center;
                        font-size:14px; color:{theme['primary']}; font-weight:700;'>
                {st.session_state.u_email[0].upper()}
            </div>
            <div style='overflow:hidden;'>
                <div style='color:#fdfaf1; font-size:12px; font-weight:600;
                            white-space:nowrap; overflow:hidden; text-overflow:ellipsis;
                            max-width:160px;'>{st.session_state.u_email}</div>
                <div style='color:rgba(253,250,241,0.4); font-size:10px;'>Tadqiqotchi</div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    current_credits = fetch_live_credits(st.session_state.u_email)
    
    # Credit display compact
    credit_percent = min((current_credits / 100) * 100, 100) if current_credits <= 100 else 100
    credit_color = "#10b981" if credit_percent > 50 else ("#f59e0b" if credit_percent > 20 else "#ef4444")
    st.markdown(f"""
        <div style='background: rgba(0,0,0,0.2); border-radius:12px; padding:14px 16px;
                    border: 1px solid rgba(197,160,89,0.1); margin-bottom:4px;'>
            <div style='display:flex; justify-content:space-between; align-items:center; margin-bottom:8px;'>
                <span style='color:rgba(253,250,241,0.5); font-size:11px; text-transform:uppercase;
                             letter-spacing:0.08em;'>💳 Kredit</span>
                <span style='color:{theme['accent']}; font-size:1.3rem; font-weight:700;
                             font-family: Playfair Display, serif;'>{current_credits}</span>
            </div>
            <div style='background:rgba(0,0,0,0.3); border-radius:6px; padding:2px; overflow:hidden;'>
                <div style='height:6px; width:{credit_percent}%; border-radius:4px;
                            background: linear-gradient(90deg, {credit_color}, {theme['accent']});
                            transition: width 0.6s ease; box-shadow: 0 0 8px {credit_color}40;'></div>
            </div>
        </div>
    """, unsafe_allow_html=True)
    
    st.divider()
    
    # === DARK MODE TOGGLE ===
    st.markdown("<p class='section-header'>🎨 Dizayn Sozlamalari</p>", unsafe_allow_html=True)
    col1, col2 = st.columns([3, 1])
    with col1:
        st.write("🌙 Tungi Rejim" if not st.session_state.dark_mode else "☀️ Kunduzgi Rejim")
    with col2:
        if st.button("🔄", key="dark_toggle"):
            st.session_state.dark_mode = not st.session_state.dark_mode
            st.rerun()
    
    # === COLOR THEME SELECTOR ===
    new_theme = st.selectbox("Rang Sxemasi", list(THEMES.keys()), key="theme_select")
    if new_theme != st.session_state.theme:
        st.session_state.theme = new_theme
        st.rerun()
    
    st.divider()
    
    # Section: Til va Xat
    st.markdown(f"<p class='section-header'>🌍 Til va Xat Tanlash</p>", unsafe_allow_html=True)
    lang = st.selectbox("Qo'lyozma tili", ["Chig'atoy", "Forscha", "Arabcha", "Eski Turkiy"])
    era = st.selectbox("Xat turi", ["Nasta'liq", "Suls", "Riq'a", "Kufiy", "Noma'lum"])
    
    st.divider()
    
    # Section: Rasm Sozlamalari
    st.markdown(f"<p class='section-header'>🎨 Rasm Sozlamalari</p>", unsafe_allow_html=True)
    br = st.slider("☀️ Yorqinlik", 0.5, 2.0, 1.0, 0.1)
    ct = st.slider("🎭 Kontrast", 0.5, 3.0, 1.3, 0.1)
    rot = st.select_slider("🔄 Aylantirish", options=[0, 90, 180, 270], value=0)

    st.divider()
    
    # === COMPARE MODE ===
    st.markdown(f"<p class='section-header'>🔄 Maxsus Rejimlar</p>", unsafe_allow_html=True)
    st.session_state.compare_mode = st.checkbox("🔄 Solishtirish Rejimi", value=st.session_state.compare_mode)
    
    st.divider()
    
    # === HISTORY SIDEBAR ===
    if st.session_state.history:
        with st.expander("📜 Tarix (Oxirgi 10 ta)"):
            for h in st.session_state.history[-10:][::-1]:
                if st.button(f"{h['date']} - {h['filename'][:25]}...", key=f"hist_{h['id']}"):
                    st.session_state.results = h['results']
                    st.session_state.chats = h.get('chats', {})
                    st.toast(f"✅ {h['filename']} yuklandi!", icon="📂")
                    st.rerun()
    
    st.divider()
    
    # KEYBOARD SHORTCUTS INFO
    with st.expander("⌨️ Klaviatura Yorliqlari"):
        st.markdown("""
            - **Esc** - Modal yopish
            - **Ctrl+Enter** - Tahlilni boshlash
            - **Ctrl+S** - Faylni saqlash
        """)
    
    st.divider()
    
    # Premium logout button
    st.markdown(f"""
        <div style='background: linear-gradient(135deg, rgba(239,68,68,0.1) 0%, rgba(220,38,38,0.15) 100%);
                    border: 1px solid rgba(239,68,68,0.2); border-radius:12px;
                    padding:14px 16px; text-align:center; margin-top:8px;
                    transition: all 0.3s ease; cursor: pointer;'
             onmouseover='this.style.background=\"linear-gradient(135deg, rgba(239,68,68,0.15) 0%, rgba(220,38,38,0.2) 100%)\"; this.style.borderColor=\"rgba(239,68,68,0.3)\";'
             onmouseout='this.style.background=\"linear-gradient(135deg, rgba(239,68,68,0.1) 0%, rgba(220,38,38,0.15) 100%)\"; this.style.borderColor=\"rgba(239,68,68,0.2)\";'>
            <span style='color:#fca5a5; font-size:12px; font-weight:600; letter-spacing:0.03em;'>🚪 TIZIMDAN CHIQISH</span>
        </div>
    """, unsafe_allow_html=True)
    
    if st.button("CHIQISH", key="logout_btn", use_container_width=True):
        st.session_state.auth = False
        st.toast("👋 Xayr!", icon="👋")
        st.rerun()

# === LANDING PAGE OR MAIN APP ===
if st.session_state.show_landing:
    render_landing_page()
    st.stop()

# === MAIN CONTENT AREA ===
st.markdown(f"""
    <div style='text-align:center; margin-bottom:36px;'>
        <h1 style='margin-bottom:8px; font-size:clamp(1.8rem, 4vw, 2.6rem);'>📜 Raqamli Qo'lyozmalar Ekspertizasi</h1>
        <p style='color:{text_secondary}; font-size:clamp(0.9rem, 2vw, 1.1rem); margin:0; line-height:1.6;'>
            Sun'iy intellekt yordamida qadimiy matnlarni tahlil qiling
        </p>
    </div>
""", unsafe_allow_html=True)

file = st.file_uploader("📤 Qo'lyozma faylini yuklang", type=["pdf", "png", "jpg", "jpeg"], label_visibility="visible")

if not file:
    # === PREMIUM EMPTY STATE ===
    st.markdown(f"""
        <div style='text-align:center; padding:80px 30px;
                    background: linear-gradient(145deg, rgba(30,42,56,0.5) 0%, rgba(22,32,44,0.7) 100%);
                    border-radius:24px; border: 2px dashed rgba(197,160,89,0.25);
                    margin: 40px 0; position: relative; overflow: hidden;'>
            <div style='position:absolute; top:0; left:0; right:0; bottom:0;
                        background: radial-gradient(circle at 50% 50%, rgba(197,160,89,0.04) 0%, transparent 70%);
                        pointer-events:none;'></div>
            <div style='font-size:4rem; margin-bottom:20px; animation: float 3s ease-in-out infinite;
                        position:relative; z-index:1;'>📜</div>
            <h3 style='color:{theme['accent']}; font-size:1.5rem; margin-bottom:12px; border:none;
                       padding:0; position:relative; z-index:1;'>Qo'lyozma yuklang</h3>
            <p style='color:rgba(253,250,241,0.5); font-size:15px; line-height:1.7;
                      max-width:400px; margin:0 auto 24px; position:relative; z-index:1;'>
                PDF, PNG, JPG yoki JPEG formatidagi fayllarni<br>yuklashingiz mumkin
            </p>
            <div style='display:flex; justify-content:center; gap:12px; flex-wrap:wrap;
                        position:relative; z-index:1;'>
                <span style='background:rgba(197,160,89,0.1); color:{theme['accent']}; padding:6px 14px;
                             border-radius:8px; font-size:0.75rem; font-weight:600;
                             border: 1px solid rgba(197,160,89,0.15);'>PDF</span>
                <span style='background:rgba(197,160,89,0.1); color:{theme['accent']}; padding:6px 14px;
                             border-radius:8px; font-size:0.75rem; font-weight:600;
                             border: 1px solid rgba(197,160,89,0.15);'>PNG</span>
                <span style='background:rgba(197,160,89,0.1); color:{theme['accent']}; padding:6px 14px;
                             border-radius:8px; font-size:0.75rem; font-weight:600;
                             border: 1px solid rgba(197,160,89,0.15);'>JPG</span>
                <span style='background:rgba(197,160,89,0.1); color:{theme['accent']}; padding:6px 14px;
                             border-radius:8px; font-size:0.75rem; font-weight:600;
                             border: 1px solid rgba(197,160,89,0.15);'>≤ 15 sahifa</span>
            </div>
        </div>
    """, unsafe_allow_html=True)
    st.stop()

if file:
    if st.session_state.get("last_fn") != file.name:
        with st.spinner("📂 Fayl tayyorlanmoqda..."):
            data = file.getvalue()
            st.session_state.file_data = data  # Fayl ma'lumotlarini saqlash
            st.session_state.is_pdf = (file.type == "application/pdf")
            
            if st.session_state.is_pdf:
                pdf = pdfium.PdfDocument(data)
                st.session_state.total_pages = min(len(pdf), 15)
                pdf.close()
                # LAZY LOAD: faqat sahifa soni hisoblanadi, rasmlar keyinroq yuklanadi
                st.session_state.imgs = [None] * st.session_state.total_pages
            else:
                st.session_state.total_pages = 1
                st.session_state.imgs = [render_page(data, 0, 1.0, False)]

            st.session_state.last_fn = file.name
            st.session_state.results, st.session_state.chats = {}, {}
            gc.collect()
            st.toast(f"✅ Fayl yuklandi! ({st.session_state.total_pages} sahifa)", icon="📁")

    # LAZY LOAD: dastlab bo'sh list, faqat tanlangan sahifalar yuklanadi
    processed = [None] * st.session_state.total_pages
    
    # Allaqachon yuklangan sahifalarni qayta ishlash (rotate/brightness/contrast)
    for page_i, im in enumerate(st.session_state.imgs):
        if im is not None:
            p = im.rotate(rot, expand=True)
            p = ImageEnhance.Brightness(p).enhance(br)
            p = ImageEnhance.Contrast(p).enhance(ct)
            processed[page_i] = p

    st.markdown("<h3 style='margin-top:30px;'>📑 Varaqlarni Tanlang</h3>", unsafe_allow_html=True)
    indices = st.multiselect(
        "Tahlil qilish uchun varaqlarni belgilang:",
        range(st.session_state.total_pages),
        default=[0],
        format_func=lambda x: f"📄 Varaq {x+1}"
    )

    # LAZY LOAD: faqat tanlangan sahifalarni yuklash
    for idx in indices:
        if idx < len(processed) and processed[idx] is None and st.session_state.is_pdf:
            with st.spinner(f"📄 Varaq {idx+1} yuklanmoqda..."):
                im = render_page(st.session_state.file_data, idx, 4.0, True)
                st.session_state.imgs[idx] = im
                p = im.rotate(rot, expand=True)
                p = ImageEnhance.Brightness(p).enhance(br)
                p = ImageEnhance.Contrast(p).enhance(ct)
                processed[idx] = p

    if not st.session_state.results and indices:
        st.markdown("<h3 style='margin-top:30px;'>🖼 Tanlangan Varaqlar</h3>", unsafe_allow_html=True)
        
        # === COMPARE MODE ===
        if st.session_state.compare_mode and len(indices) >= 2:
            st.info("🔄 Solishtirish rejimi: Birinchi 2 ta varaq yonma-yon")
            c1, c2 = st.columns(2)
            with c1:
                st.markdown(f"<h4>📄 Varaq {indices[0]+1}</h4>", unsafe_allow_html=True)
                st.markdown('<div class="magnifier-container">', unsafe_allow_html=True)
                st.image(processed[indices[0]], use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)
            with c2:
                st.markdown(f"<h4>📄 Varaq {indices[1]+1}</h4>", unsafe_allow_html=True)
                st.markdown('<div class="magnifier-container">', unsafe_allow_html=True)
                st.image(processed[indices[1]], use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)
        else:
            cols = st.columns(min(len(indices), 3))
            for i, idx in enumerate(indices):
                with cols[i % 3]:
                    st.markdown('<div class="magnifier-container">', unsafe_allow_html=True)
                    st.image(processed[idx], caption=f"Varaq {idx+1}", use_container_width=True)
                    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    
    if st.button("✨ TAHLILNI BOSHLASH"):
        if current_credits >= len(indices):
            prompt = f"""Sen qadimiy qo'lyozmalar bo'yicha DUNYO DARAJASIDAGI EKSPERT PALEOGRAF va FILOLOG sifatida ish ko'r.

═══════════════════════════════════════════
📥 KIRISH MA'LUMOTLARI:
═══════════════════════════════════════════
- Foydalanuvchi ko'rsatgan til: {lang} (faqat yo'naltiruvchi, haqiqiy tilni RASMDAN aniqla)
- Foydalanuvchi ko'rsatgan xat turi: {era} (faqat yo'naltiruvchi, haqiqiy xatni RASMDAN aniqla)

⚠️ ASOSIY VAZIFA: Rasmni DIQQAT BILAN o'rgan va MAKSIMAL ANIQLIK bilan tahlil qil.

═══════════════════════════════════════════
📋 MAJBURIY TAHLIL STRUKTURASI (7 BO'LIM):
═══════════════════════════════════════════

## 1. MANBA IDENTIFIKATSIYASI
- **Til**: [RASMNI o'rganib ANIQ aniqlang: Chig'atoy / Forscha / Arabcha / Eski Turkiy / Aralash]
- **Xat turi**: [Xat uslubini ANIQ aniqlang: Nasta'liq / Naskh / Suls / Shikasta / Riq'a / Kufiy / Boshqa]
- **Taxminiy davr**: [Asrni aniqlang va ASOSLANG, masalan: "XV-XVI asr, chunki..."]
- **O'qilish darajasi**: [Rasmiy baho: Yaxshi o'qiladi (80-100%) / Qisman o'qiladi (50-79%) / Qiyin o'qiladi (0-49%)]
- **Qog'oz holati**: [Yaxshi saqlangan / Qisman shikastlangan / Jiddiy zarar ko'rgan]

## 2. TRANSLITERATSIYA (Asl Arab-Fors yozuvi)
Rasmdagi HAR BIR qatorni RAQAM bilan belgilab, ARAB/FORS ALIFBOSIDA yoz:

[1] بسم الله الرحمن الرحيم
[2] الحمد لله رب العالمين
[3] ...

🔍 QOIDALAR:
- MAJBURIY: Har bir qatorni [1], [2], [3]... raqam bilan BELGILAB yoz
- Rasmdagi FIZIK qator tartibini AYNAN saqlang — aralashtirib YOZISH TAQIQLANGAN
- Noaniq harflar uchun: ن[?-nuqta ko'rinmaydi] yoki ک[?-shikastlangan]
- Yo'qolgan/o'chirilgan: [...5-6 so'z yo'qolgan] yoki [1 satr o'chirilgan]
- Harakat belgilarini (zabar, zer, pesh) ALBATTA ko'rsating
- Kamida 3-5 qator bo'lishi SHART

## 3. LOTIN TRANSKRIPSIYASI
Yuqoridagi AYNI raqamlar bilan lotin alifbosida yoz (raqamlar MOS kelishi SHART):

[1] Bismillahir rohmanir rohiym
[2] Alhamdulillahi robbil 'alamin
[3] ...

📝 STANDART: O'zbek lotin alifbosi
- Asosiy: a, b, d, e, f, g, gʻ, h, i, j, k, l, m, n, o, p, q, r, s, sh, t, u, v, x, y, z, oʻ, ch, ng
- Qo'shimcha: ā (uzun a), ī (uzun i), ū (uzun u), ʼ (hamza), ʻ (ayn)
- HAR BIR SO'ZNI transliteratsiya qiling
- QATOR RAQAMI transliteratsiya bilan AYNAN MOS kelishi SHART

## 4. TO'LIQ TARJIMA (Zamonaviy o'zbek tilida)
[Butun matnni ZAMONAVIY O'ZBEK TILIGA to'liq tarjima qiling]

📖 QOIDALAR:
- Tarjima MAZMUNNI TO'LIQ aks ettirishi SHART
- Adabiy, tushunarli, ravon til
- Qadimiy iboralarni zamonaviy shaklga o'tkaz
- Noaniq joylar: [taxminan: men shunday o'qidim] shaklida
- Kamida 3-5 gap bo'lishi SHART

## 5. LEKSIK-SEMANTIK TAHLIL
[Qadimiy, kam uchraydigan, muhim so'zlarni JADVALDA ko'rsat]

| № | Asl so'z (Arab) | Lotin | Ma'nosi | Zamonaviy ekvivalent | Izoh |
|---|-----------------|-------|---------|---------------------|------|
| 1 | مثال | misal | namuna | misol | umumiy so'z |
| 2 | ... | ... | ... | ... | ... |

⚠️ KAMIDA 5 TA, IDEAL 8-10 TA SO'Z ko'rsating.

## 6. AKADEMIK IZOHLAR
- **Paleografik xususiyatlar**: [Xat uslubi tavsifi, qalam kengligi, siyoh turi/rangi, yozuv yo'nalishi]
- **Tarixiy kontekst**: [Qaysi davr, qaysi hudud/shahar, siyosiy vaziyat, madaniy muhit]
- **Til xususiyatlari**: [Arxaik so'zlar, grammatik shakllar, morfologiya, sintaksis]
- **Mazmun tahlili**: [Matn mavzusi, janri (xat/hujjat/she'r/ilmiy), ahamiyati, muallif uslubi]
- **Qiyinchiliklar**: [Aniq nima qiyin bo'ldi: siyoh o'chgan, qog'oz yirtilgan, xat noaniq...]

## 7. ANIQLIK BAHOSI (ANIQ FOIZLAR)
| Mezon | Foiz | Batafsil Izoh |
|-------|------|---------------|
| Transliteratsiya aniqligi | X% | Nechta harf/so'z aniq, nechta [?] |
| Tarjima aniqligi | X% | Mazmun tushunilish darajasi |
| Leksik tahlil to'liqligi | X% | So'zlar soni va tahlil chuqurligi |
| Umumiy ishonch darajasi | X% | Umumiy natijaga ishonch |

🚨 **Asosiy qiyinchiliklar ro'yxati**:
1. [Birinchi muammo va uning sababi]
2. [Ikkinchi muammo]
3. [Uchinchi muammo]

═══════════════════════════════════════════
⚠️ MUHIM QOIDALAR (MAJBURIY BAJARISH):
═══════════════════════════════════════════
1. ✅ Barcha javoblar O'ZBEK TILIDA
2. ✅ BARCHA 7 BO'LIM to'ldirilishi MAJBURIY - bo'sh bo'lim QO'YISH MUMKIN EMAS
3. ✅ Noaniqliklarni [?] bilan belgilab, SABAB yozing
4. ✅ Bu oddiy tarjima EMAS - bu PROFESSIONAL ILMIY EKSPERTIZA
5. ✅ Jadvalda KAMIDA 5 TA qator bo'lishi SHART
6. ✅ Har bir bo'limda KAMIDA 3-5 qator matn bo'lishi kerak
7. ✅ Aniqlik foizlarini REAL baholang (75-95% oralig'ida bo'lishi kerak)
8. ❌ "O'qib bo'lmaydi", "Ko'rinmaydi", "Yo'q" kabi umumiy javoblar TAQIQLANGAN
9. ✅ Agar matn juda qiyin bo'lsa - iloji boricha o'qib, [?] bilan belgilang
10. ✅ Har bir [?] belgisi uchun SABAB yozing (masalan: [?-siyoh o'chgan])

💡 ESLATMA: Sifatli javob = Ko'proq ma'lumot + Kam [?] belgilar + To'liq jadval + Aniq foizlar"""
            
            # === PREMIUM PROGRESS TRACKER ===
            st.markdown(f"""
                <div style='background: rgba(30,42,56,0.6); border-radius:14px; padding:20px 24px;
                            border: 1px solid rgba(197,160,89,0.12); margin:24px 0;'>
                    <div style='display:flex; align-items:center; gap:10px; margin-bottom:10px;'>
                        <span style='font-size:20px;'>🔍</span>
                        <span style='color:{theme['accent']}; font-weight:600; font-size:15px;'>Tahlil jarayoni</span>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            progress_bar = st.progress(0, text="🔍 Tahlil boshlanmoqda...")
            total_attempts = 0
            quality_issues = []
            
            for i, idx in enumerate(indices):
                with st.status(f"🔍 Varaq {idx+1} ekspertizadan o'tkazilmoqda...") as status:
                    try:
                        current_img = processed[idx]
                        
                        # === SMART CROPPING CHECK ===
                        if should_split_image(current_img):
                            status.update(label=f"📐 Varaq {idx+1} bo'laklarga bo'linmoqda...")
                            crops = split_image_smart(current_img)
                            crop_results = []
                            
                            for j, crop in enumerate(crops):
                                status.update(label=f"🔍 Varaq {idx+1} - Qism {j+1}/{len(crops)}...")
                                # Har bir qismga ANIQ ko'rsatma berish
                                crop_prompt = prompt + f"\n\n⚠️ DIQQAT: Bu rasmning {j+1}-QISMI ({len(crops)} qismdan). Faqat SHU rasmdagi matnni o'qi. Oldingi/keyingi qismlar alohida tahlil qilinadi."
                                result, quality, attempts = analyze_with_retry(model, crop_prompt, crop, max_retries=1)
                                total_attempts += attempts
                                
                                if result:
                                    crop_results.append(result)
                                    if quality["score"] < 70:
                                        quality_issues.append(f"Varaq {idx+1} qism {j+1}: {quality['reason']}")
                            
                            if crop_results:
                                st.session_state.results[idx] = merge_results(crop_results)
                                use_credit_atomic(st.session_state.u_email)
                                st.toast(f"✅ Varaq {idx+1} tayyor! ({len(crops)} qism)", icon="🎉")
                                st.success(f"✅ Varaq {idx+1} muvaffaqiyatli tahlil qilindi")
                            else:
                                st.error(f"⚠️ Varaq {idx+1} tahlil qilinmadi")
                        
                        else:
                            # === NORMAL ANALYSIS WITH RETRY ===
                            result, quality, attempts = analyze_with_retry(model, prompt, current_img, max_retries=2)
                            total_attempts += attempts
                            
                            if result:
                                st.session_state.results[idx] = result
                                use_credit_atomic(st.session_state.u_email)
                                
                                # Quality indicator
                                # Sifat bo'yicha xabar berish (YANGI CHEGARALAR)
                                if quality["score"] >= 85:
                                    st.toast(f"🏆 Varaq {idx+1} - A'lo sifat!", icon="🎉")
                                elif quality["score"] >= 70:
                                    st.toast(f"✅ Varaq {idx+1} - Yaxshi sifat", icon="✅")
                                elif quality["score"] >= 55:
                                    st.toast(f"⚠️ Varaq {idx+1} - Qoniqarli sifat", icon="⚠️")
                                    quality_issues.append(f"Varaq {idx+1}: {quality['reason']}")
                                else:
                                    st.toast(f"❌ Varaq {idx+1} - Past sifat", icon="❌")
                                    quality_issues.append(f"Varaq {idx+1}: {quality['reason']}")
                                
                                # YANGI: Batafsil sifat hisobotini session_state'ga saqlash
                                if 'quality_reports' not in st.session_state:
                                    st.session_state.quality_reports = {}
                                st.session_state.quality_reports[idx] = quality
                                
                                if attempts > 1:
                                    st.info(f"ℹ️ {attempts} urinishda tahlil qilindi")
                                st.success(f"✅ Varaq {idx+1} muvaffaqiyatli tahlil qilindi")
                            else:
                                st.error(f"⚠️ Varaq {idx+1}: AI javob berish imkoniyatiga ega emas")
                                
                    except Exception as e:
                        st.error(f"❌ Xatolik yuz berdi: {e}")
                
                # Update progress with custom styling
                progress_percent = (i+1)/len(indices)
                progress_bar.progress(progress_percent, text=f"📊 {i+1}/{len(indices)} varaq tahlil qilindi ({int(progress_percent*100)}%)")
            
            # === QUALITY SUMMARY ===
            if quality_issues:
                with st.expander("⚠️ Sifat ogohlantirishlari", expanded=False):
                    for issue in quality_issues:
                        st.warning(issue)
                    st.info("💡 Tavsiya: Rasmlar sifatini oshiring yoki aniqroq skanerdan foydalaning")
            
            # Save to history
            st.session_state.history.append({
                'id': datetime.now().timestamp(),
                'date': datetime.now().strftime("%d.%m.%Y %H:%M"),
                'filename': file.name,
                'results': st.session_state.results.copy(),
                'chats': st.session_state.chats.copy(),
                'quality_issues': quality_issues,
                'total_attempts': total_attempts
            })
            
            # Keep only last 10
            if len(st.session_state.history) > 10:
                st.session_state.history = st.session_state.history[-10:]
            
            progress_bar.empty()
            st.balloons()
            st.toast("🎉 Barcha varaqlar tahlil qilindi!", icon="✨")
            st.rerun()
        else:
            st.warning(f"⚠️ Kredit yetarli emas! Sizda {current_credits} sahifa kredit bor, {len(indices)} sahifa tahlil qilish uchun yetarli emas.")

    if st.session_state.results:
        st.divider()
        
        # === PROFESSIONAL HEADER ===
        result_count = len(st.session_state.results)
        st.markdown(f"""
        <div style='background: linear-gradient(145deg, rgba(12,20,33,0.98) 0%, rgba(22,32,44,0.99) 100%);
                    border: 1px solid rgba(197,160,89,0.2);
                    border-radius: 20px;
                    padding: 28px 32px;
                    margin-bottom: 28px;
                    box-shadow: 0 12px 40px rgba(0,0,0,0.25), 0 0 40px rgba(197,160,89,0.04);
                    position: relative; overflow: hidden;
                    animation: fadeInUp 0.6s cubic-bezier(0.23, 1, 0.32, 1);'>
            <div style='position:absolute; top:0; left:0; right:0; height:2px;
                        background: linear-gradient(90deg, {theme['accent']}, {theme['accent2']}, {theme['accent']});
                        background-size: 200% 100%; animation: shimmer 3s ease-in-out infinite;'></div>
            <div style='display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:16px;'>
                <div>
                    <h2 style='margin:0; padding:0; border:none; color:{theme['accent']};
                               font-family:Playfair Display,serif; font-size:1.5rem; text-align:left;'>
                        📊 Akademik Ekspertiza Natijasi
                    </h2>
                    <p style='margin:6px 0 0; color:rgba(253,250,241,0.45); font-size:13px;'>
                        {result_count} varaq tahlil qilindi · {file.name}
                    </p>
                </div>
                <div style='display:flex; gap:8px; flex-wrap:wrap;'>
                    <span style='background: linear-gradient(135deg, {theme['accent']}, {theme['accent2']});
                                 color:{theme['primary']}; padding:5px 14px; border-radius:20px;
                                 font-size:11px; font-weight:700; letter-spacing:0.03em;'>✅ TEKSHIRILGAN</span>
                    <span style='background:rgba(16,185,129,0.12); color:#10b981;
                                 padding:5px 14px; border-radius:20px; font-size:11px; font-weight:600;
                                 border:1px solid rgba(16,185,129,0.2);'>🔬 AKADEMIK</span>
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # === EXPORT FORMAT SELECTOR ===
        st.markdown(f"""
            <div style='background: rgba(30,42,56,0.6); border-radius:14px; padding:20px 24px;
                        border: 1px solid rgba(197,160,89,0.12); margin-bottom:24px;
                        display:flex; align-items:center; justify-content:space-between; flex-wrap:wrap; gap:12px;'>
                <div style='display:flex; align-items:center; gap:8px;'>
                    <span style='font-size:20px;'>📥</span>
                    <span style='color:rgba(253,250,241,0.6); font-size:14px;'>Natijalarni yuklab oling</span>
                </div>
            </div>
        """, unsafe_allow_html=True)
        
        exp_c1, exp_c2 = st.columns([3, 1])
        with exp_c2:
            export_format = st.selectbox("📥 Format", ["DOCX", "TXT", "JSON"], label_visibility="collapsed")
        
        final_text = ""
        today = datetime.now().strftime("%d.%m.%Y")
        
        # Mobile navigation buttons
        result_indices = sorted(st.session_state.results.keys())
        if len(result_indices) > 1:
            nav_col1, nav_col2, nav_col3 = st.columns([1, 2, 1])
            with nav_col1:
                if st.button("⬅️ Oldingi", key="mobilePrev", use_container_width=True):
                    if st.session_state.current_page_index > 0:
                        st.session_state.current_page_index -= 1
                        st.rerun()
            with nav_col3:
                if st.button("Keyingi ➡️", key="mobileNext", use_container_width=True):
                    if st.session_state.current_page_index < len(result_indices) - 1:
                        st.session_state.current_page_index += 1
                        st.rerun()

        for idx in result_indices:
            st.markdown(f"""
                <div style='margin-top:40px; display:flex; align-items:center; gap:12px; margin-bottom:16px;'>
                    <div style='background: linear-gradient(135deg, {theme['accent']}, {theme['accent2']});
                                width:36px; height:36px; border-radius:10px;
                                display:flex; align-items:center; justify-content:center;
                                font-size:16px; font-weight:700; color:{theme['primary']};
                                box-shadow: 0 4px 12px rgba(197,160,89,0.25);'>{idx+1}</div>
                    <h4 style='margin:0; padding:0; border:none; text-align:left; font-size:1.2rem;'>📖 Varaq {idx+1} - Tahlil</h4>
                </div>
            """, unsafe_allow_html=True)
            c1, c2 = st.columns([1, 1.3])

            with c1:
                st.markdown('<div class="magnifier-container">', unsafe_allow_html=True)
                st.image(processed[idx], use_container_width=True)
                st.markdown("</div>", unsafe_allow_html=True)

            with c2:
                # YANGI: Sifat hisobotini ko'rsatish
                if hasattr(st.session_state, 'quality_reports') and idx in st.session_state.quality_reports:
                    quality_report_html = generate_quality_report(st.session_state.quality_reports[idx], theme)
                    st.markdown(quality_report_html, unsafe_allow_html=True)
                
                # Natijani markdown sifatida ko'rsatish
                st.markdown("<div class='result-box'>", unsafe_allow_html=True)
                st.markdown(st.session_state.results[idx])
                st.markdown("</div>", unsafe_allow_html=True)
                
                cite = f"Iqtibos: Manuscript AI (2026). Varaq {idx+1} tahlili ({lang}). Ekspert: d87809889-dot. Sana: {today}."
                st.markdown(f"<div class='citation-box'>📌 {cite}</div>", unsafe_allow_html=True)

                st.markdown(f"<p style='color:{text_secondary}; font-weight:bold; margin-top:20px; margin-bottom:8px;'>✏️ Tahrirlash:</p>", unsafe_allow_html=True)
                st.session_state.results[idx] = st.text_area(
                    f"Natijani tahrirlash",
                    value=st.session_state.results[idx],
                    height=350,
                    key=f"ed_{idx}",
                    label_visibility="collapsed"
                )

                final_text += f"\n\n--- PAGE {idx+1} ---\n{st.session_state.results[idx]}\n\n{cite}"

                # === CHAT INTERFACE ===
                st.markdown(f"""
                    <div style='margin-top:25px; margin-bottom:12px; display:flex; align-items:center; gap:8px;'>
                        <span style='font-size:18px;'>💬</span>
                        <span style='color:{text_secondary}; font-weight:600; font-size:14px;'>Savollar va Javoblar</span>
                    </div>
                """, unsafe_allow_html=True)
                
                st.session_state.chats.setdefault(idx, [])
                for ch in st.session_state.chats[idx]:
                    st.markdown(f"<div class='chat-user'><b>Savol:</b> {ch['q']}</div>", unsafe_allow_html=True)
                    st.markdown(f"<div class='chat-ai'><b>Javob:</b> {ch['a']}</div>", unsafe_allow_html=True)

                q = st.text_input("🤔 Savolingizni yozing:", key=f"q_in_{idx}", placeholder="Matn haqida savol bering...")
                if st.button(f"📤 So'rash", key=f"btn_{idx}"):
                    if q:
                        with st.spinner("🤖 AI javob tayyorlayapti..."):
                            chat_prompt = f"""Sen qadimiy qo'lyozmalar bo'yicha EKSPERT sifatida javob ber.

═══════════════════════════════════════════
KONTEKST:
═══════════════════════════════════════════
Bu qo'lyozma hujjati haqida oldingi tahlil qilingan:

{st.session_state.results[idx]}

═══════════════════════════════════════════
FOYDALANUVCHI SAVOLI:
═══════════════════════════════════════════
{q}

═══════════════════════════════════════════
JAVOB QOIDALARI:
═══════════════════════════════════════════
1. Javobni O'ZBEK TILIDA ber
2. Aniq va to'liq javob ber (1-3 paragraf)
3. Agar javob tahlilda bo'lsa - aniq ko'rsat va iqtibos keltir
4. Agar javob tahlilda YO'Q bo'lsa - RASMNI qayta ko'rib javob ber
5. Agar umuman javob berib bo'lmasa - sababini aniq tushuntir
6. Taxminiy javob bo'lsa - [TAXMIN] deb belgilab ber"""
                            chat_res = model.generate_content([
                                chat_prompt,
                                img_to_png_payload(processed[idx])
                            ])
                            st.session_state.chats[idx].append({"q": q, "a": chat_res.text})
                            st.toast("✅ Javob olindi!", icon="💬")
                            st.rerun()
                    else:
                        st.warning("⚠️ Iltimos, avval savol yozing!")

        if final_text:
            st.divider()
            
            # === EXPORT BASED ON FORMAT ===
            if export_format == "DOCX":
                doc = Document()
                doc.add_paragraph(final_text)
                bio = io.BytesIO()
                doc.save(bio)
                st.download_button(
                    "📥 WORD FORMATDA YUKLAB OLISH",
                    bio.getvalue(),
                    "manuscript_ai_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    on_click=lambda: st.toast("✅ DOCX yuklab olindi!", icon="📥")
                )
            elif export_format == "TXT":
                st.download_button(
                    "📥 TEXT FORMATDA YUKLAB OLISH",
                    final_text,
                    "manuscript_ai_report.txt",
                    mime="text/plain",
                    use_container_width=True,
                    on_click=lambda: st.toast("✅ TXT yuklab olindi!", icon="📥")
                )
            elif export_format == "JSON":
                json_data = json.dumps({
                    "metadata": {
                        "date": today,
                        "language": lang,
                        "script": era,
                        "filename": file.name,
                        "user": st.session_state.u_email
                    },
                    "results": st.session_state.results,
                    "chats": st.session_state.chats
                }, ensure_ascii=False, indent=2)
                st.download_button(
                    "📥 JSON FORMATDA YUKLAB OLISH",
                    json_data,
                    "manuscript_ai_report.json",
                    mime="application/json",
                    use_container_width=True,
                    on_click=lambda: st.toast("✅ JSON yuklab olindi!", icon="📥")
                )

gc.collect()
